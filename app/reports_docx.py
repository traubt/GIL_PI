# app/reports_docx.py
import io
import json
import os
import datetime
from flask import Blueprint, request, send_file, abort, jsonify, render_template_string, current_app, url_for
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
import mammoth
import os, io, datetime
import subprocess, tempfile, shutil, time, glob
from pathlib import Path
from .models import *
from typing import List

# ---- Use app config for paths ----
from .config import Config
import json
from datetime import datetime as _dt
from sqlalchemy import func
from flask import session


DOCX_TEMPLATES_DIR = getattr(
    Config, "DOCX_TEMPLATES_DIR",
    os.environ.get("DOCX_TEMPLATES_DIR", "/mnt/data")
)
OUTPUT_DIR = getattr(
    Config, "DOCX_OUTPUT_DIR",
    os.environ.get("DOCX_OUTPUT_DIR", "/tmp/gil_reports")
)

reports_docx_bp = Blueprint("reports_docx", __name__, url_prefix="/reports")

def _iso_to_ddmmyyyy_dash(s: str) -> str:
    """
    '2026-01-27' -> '27-01-2026'
    Leaves unknown formats as-is.
    """
    if not s:
        return ""
    s = str(s).strip()
    try:
        y, m, d = s.split("-")
        if len(y) == 4 and len(m) == 2 and len(d) == 2:
            return f"{d}-{m}-{y}"
    except Exception:
        pass
    return s



# ============================================================
# API: Investigator Tracking Reports (per-insured, per-day)
# - Dor requirement: 1 activity report per day per insured.
# - We expose dates + activities by report_date (not report_id).
# ============================================================
@reports_docx_bp.route("/api/insured/<int:insured_id>/tracking-dates", methods=["GET"])
def api_tracking_dates(insured_id: int):
    from datetime import date as _date
    try:
        rows = (
            GilTrackingReport.query
            .filter(GilTrackingReport.insured_id == insured_id)
            .order_by(GilTrackingReport.report_date.desc())
            .all()
        )
        seen = set()
        out = []
        for r in rows:
            if not r.report_date:
                continue
            iso = r.report_date.isoformat()
            if iso in seen:
                continue
            seen.add(iso)
            dmy = r.report_date.strftime("%d/%m/%Y")
            out.append({"iso": iso, "dmy": dmy})
        return jsonify({"status": "ok", "dates": out})
    except Exception as e:
        current_app.logger.exception("api_tracking_dates failed")
        return jsonify({"status": "error", "error": str(e)}), 500


@reports_docx_bp.route("/api/insured/<int:insured_id>/tracking-activities", methods=["GET"])
def api_tracking_activities(insured_id: int):
    from datetime import datetime as _dt
    from sqlalchemy import func

    report_date = (request.args.get("report_date") or "").strip()

    # Parse YYYY-MM-DD
    try:
        dt = _dt.strptime(report_date, "%Y-%m-%d").date()
    except Exception:
        return jsonify({"status": "ok", "activities": []})

    try:
        rep_row = (
            GilTrackingReport.query
            .filter(
                GilTrackingReport.insured_id == insured_id,
                GilTrackingReport.report_date == dt
            )
            .first()
        )
        if not rep_row:
            return jsonify({"status": "ok", "activities": []})

        # --------------------------------------------------
        # 1) Prefer CURRENT set (admin override if exists)
        # --------------------------------------------------
        acts = (
            GilTrackingReportActivity.query
            .filter(
                GilTrackingReportActivity.report_id == rep_row.report_id,
                GilTrackingReportActivity.is_current == True
            )
            .order_by(
                GilTrackingReportActivity.sort_order.asc(),
                GilTrackingReportActivity.activity_id.asc(),
            )
            .all()
        )

        # --------------------------------------------------
        # 2) Fallback: latest set_no (safety net)
        # --------------------------------------------------
        if not acts:
            max_set = (
                db.session.query(func.max(GilTrackingReportActivity.set_no))
                .filter(GilTrackingReportActivity.report_id == rep_row.report_id)
                .scalar()
            )

            if max_set is not None:
                acts = (
                    GilTrackingReportActivity.query
                    .filter(
                        GilTrackingReportActivity.report_id == rep_row.report_id,
                        GilTrackingReportActivity.set_no == max_set
                    )
                    .order_by(
                        GilTrackingReportActivity.sort_order.asc(),
                        GilTrackingReportActivity.activity_id.asc(),
                    )
                    .all()
                )

        out = []
        for a in acts:
            t = a.activity_time.strftime("%H:%M") if a.activity_time else ""
            out.append({
                "time": t,
                "description": a.description or "",
                # optional metadata (harmless to UI)
                "source": getattr(a, "source", "investigator"),
                "set_no": getattr(a, "set_no", 1),
                "is_current": getattr(a, "is_current", True),
            })

        return jsonify({"status": "ok", "activities": out})

    except Exception as e:
        current_app.logger.exception("api_tracking_activities failed")
        return jsonify({"status": "error", "error": str(e)}), 500


@reports_docx_bp.route("/api/insured/<int:insured_id>/tracking-activities/save", methods=["POST"])
def api_tracking_activities_save(insured_id: int):

    # Auth (match your app style)
    user_data = session.get("user")
    user = json.loads(user_data) if user_data else {}
    if not user:
        return jsonify({"status": "error", "message": "Not logged in"}), 401

    payload = request.get_json(silent=True) or {}
    report_date = (payload.get("report_date") or "").strip()
    activities  = payload.get("activities") or []

    if not report_date:
        return jsonify({"status": "error", "message": "Missing report_date"}), 400
    if not isinstance(activities, list):
        return jsonify({"status": "error", "message": "Invalid activities"}), 400

    # Parse YYYY-MM-DD
    try:
        dt = _dt.strptime(report_date, "%Y-%m-%d").date()
    except Exception:
        return jsonify({"status": "error", "message": "Invalid report_date format"}), 400

    rep_row = (
        GilTrackingReport.query
        .filter(GilTrackingReport.insured_id == insured_id,
                GilTrackingReport.report_date == dt)
        .first()
    )
    if not rep_row:
        return jsonify({"status": "error", "message": "Tracking report not found"}), 404

    # Normalize + validate
    cleaned = []
    for idx, a in enumerate(activities):
        t = (a.get("time") or "").strip()
        d = (a.get("description") or "").strip()

        if not t or not d:
            continue

        # accept HH:MM or HH:MM:SS
        try:
            if len(t) == 5:
                tt = _dt.strptime(t, "%H:%M").time()
            else:
                tt = _dt.strptime(t, "%H:%M:%S").time()
        except Exception:
            return jsonify({"status": "error", "message": f"Invalid time: {t}"}), 400

        cleaned.append({"time": tt, "description": d, "sort_order": idx})

    if len(cleaned) < 2:
        return jsonify({"status": "error", "message": "Need at least start and end activities"}), 400

    user_id = user.get("id")

    try:
        max_set = (
            db.session.query(func.max(GilTrackingReportActivity.set_no))
            .filter(GilTrackingReportActivity.report_id == rep_row.report_id)
            .scalar()
        ) or 1
        new_set = int(max_set) + 1

        # flip current -> non-current
        (GilTrackingReportActivity.query
         .filter(GilTrackingReportActivity.report_id == rep_row.report_id,
                 GilTrackingReportActivity.is_current == True)
         .update({"is_current": False}, synchronize_session=False))

        # insert new current set
        for r in cleaned:
            db.session.add(GilTrackingReportActivity(
                report_id=rep_row.report_id,
                set_no=new_set,
                is_current=True,
                source="admin",
                created_by_user_id=user_id,
                activity_time=r["time"],
                description=r["description"],
                sort_order=r["sort_order"],
            ))

        db.session.commit()
        return jsonify({"status": "ok", "message": "Activities saved", "set_no": new_set})

    except Exception as e:
        db.session.rollback()
        current_app.logger.exception("api_tracking_activities_save failed")
        return jsonify({"status": "error", "message": str(e)}), 500



TEMPLATE_MAP = {
    "tracking": "menora_report.docx",   # your old tracking template
    "siudi":    "menora_siudi.docx",    # the new one
     "siudi_invoice": "invoice_monora_siudi.docx",

    # === new Menora Life insurance templates ===
    "menora_life_followup": "menora_life_followup.docx",
    "menora_life_photos": "menora_photos.docx",  # placeholder until you build it
    "menora_life_photoid": "menora_photo_id.docx",  # reuse Siudi photo ID
    "menora_life_invoice": "invoice_menora_life.docx",  # placeholder until ready
}

from datetime import datetime, date
from flask import request, current_app

from dataclasses import dataclass

@dataclass
class PreparedPhoto:
    path: str
    orientation: str   # "portrait" | "landscape"
    image: InlineImage


def ddmmyyyy(iso_str: str) -> str:
    try:
        return datetime.strptime(iso_str.strip(), "%Y-%m-%d").strftime("%d/%m/%Y")
    except Exception:
        return iso_str or ""

def map_template_key(key: str) -> str:
    return TEMPLATE_MAP.get((key or "").lower(), TEMPLATE_MAP["tracking"])


def _deep_set(d: dict, dotted: str, value):
    """set nested dict value by dotted path (e.g., 'insured.name')."""
    parts = dotted.split(".")
    cur = d
    for p in parts[:-1]:
        if p not in cur or not isinstance(cur[p], dict):
            cur[p] = {}
        cur = cur[p]
    cur[parts[-1]] = value
    return d

def pick_template_for_report(report_id: int) -> str:
    # Phase-1: always use the Menora 'Siudi' template
    return "menora_siudi.docx"

def _parse_tracking_rows(raw: str):
    """
    Parse lines like '06:00 - הגיע לכתובת...' into
    [{'time': '06:00', 'text': 'הגיע לכתובת...'}, ...]
    """
    rows = []
    if not raw:
        return rows

    for line in raw.splitlines():
        line = (line or "").strip()
        if not line:
            continue

        time_part = ""
        text_part = ""

        # Accept 'HH:MM - text' or 'HH:MM text'
        # First try a dash close to the start
        dash_idx = line.find("-")
        if 0 <= dash_idx <= 6:
            time_part = line[:dash_idx].strip()
            text_part = line[dash_idx + 1 :].strip()
        else:
            parts = line.split(None, 1)
            if len(parts) == 2:
                time_part, text_part = parts[0].strip(), parts[1].strip()
            else:
                # no clear time, treat entire line as text
                text_part = line

        rows.append({"time": time_part, "text": text_part})

    return rows

# ---- context builder (KEEP THIS VERSION) ----
from datetime import datetime, date

def _fmt_d(d: date | None, fmt="%d/%m/%Y") -> str:
    return d.strftime(fmt) if d else ""

def _calc_age(birth: date | None, on: date | None = None) -> int | str:
    if not birth:
        return ""
    on = on or date.today()
    years = on.year - birth.year - ((on.month, on.day) < (birth.month, birth.day))
    return years

def _iso_to_dots(iso: str | None) -> str:
    """
    'YYYY-MM-DD' -> 'D.MM.YYYY'  (no leading zero on day; month stays 2-digit)
    """
    if not iso or len(iso) < 10:
        return ""
    y, m, d = iso[:10].split("-")
    # remove leading zero from day
    d = str(int(d))
    return f"{d}.{m}.{y}"

from datetime import datetime

def _now_hebrew() -> str:
    """Return date like: '2025 בנובמבר 4' (Hebrew month with ב- prefix, no leading zero)."""
    months = ["ינואר","פברואר","מרץ","אפריל","מאי","יוני",
              "יולי","אוגוסט","ספטמבר","אוקטובר","נובמבר","דצמבר"]
    dt = datetime.now()
    month_he = f"ב{months[dt.month - 1]}"
    return f"{dt.day} {month_he} {dt.year}"


import re

def _to_iso(date_str: str | None) -> str:
    """
    Accept 'YYYY-MM-DD' or 'DD/MM/YYYY' (also D/M/YYYY) and return ISO 'YYYY-MM-DD'.
    Returns '' if parsing fails.
    """
    if not date_str:
        return ""
    s = date_str.strip()
    # Already ISO?
    if re.fullmatch(r"\d{4}-\d{2}-\d{2}", s):
        return s
    # DD/MM/YYYY or D/M/YYYY
    m = re.fullmatch(r"(\d{1,2})/(\d{1,2})/(\d{4})", s)
    if m:
        d, mth, y = m.groups()
        return f"{y}-{int(mth):02d}-{int(d):02d}"
    return ""  # unknown format

def _iso_to_dots(iso: str | None) -> str:
    """
    Accept ISO 'YYYY-MM-DD' (or DD/MM/YYYY defensively) and return 'D.MM.YYYY'.
    """
    if not iso:
        return ""
    # normalize if got DD/MM/YYYY by mistake
    if "/" in iso:
        iso = _to_iso(iso)
    if not iso or len(iso) < 10 or "-" not in iso:
        return ""
    y, m, d = iso[:10].split("-")
    d = str(int(d))             # no leading zero for day
    return f"{d}.{m}.{y}"

# put near other helpers
def _gender_lex(gender_raw: str | None) -> dict:
    """
    Returns Hebrew words that change by gender.
    Accepts values like 'זכר'/'נקבה' or 'male'/'female'/'M'/'F'.
    """
    g = (gender_raw or "").strip().lower()
    is_female = g in {"נקבה", "female", "f", "נשים", "אשה", "אישה"}
    return {
        "is_female": is_female,
        "insured_label": "מבוטחת" if is_female else "מבוטח",
        "born_word":     "ילידת"   if is_female else "יליד",
        "claims_verb":   "טוענת"   if is_female else "טוען",
        "obj_suffix":    "ה"        if is_female else "ו",  # לתעדה / לתעדו
        "house_poss":    "ה"        if is_female else "ו",  # ביתה / ביתו
        "state_poss":    "ה"        if is_female else "ו",  # מצבה / מצבו
        "func_poss":     "ה"        if is_female else "ו",  # תפקודה / תפקודו
        "photos_poss":   "תמונותיה" if is_female else "תמונותיו",
        "pronoun_subj":  "היא"      if is_female else "הוא",
        "pronoun_obj":   "אותה"     if is_female else "אותו",
        "injured": "הנפגעת" if is_female else "הנפגע",
        "work": "עובדת" if is_female else "עובד",
        "functioning": "תפקודה" if is_female else "תפקודו",

    }



def _fetch_insured_row(insured_id: int) -> dict:
    ins = GilInsured.query.get(insured_id) if insured_id else None
    if not ins:
        return {}
    full_name = (" ".join(filter(None, [ins.last_name, ins.first_name]))).strip()
    birth_year = ins.birth_date.year if getattr(ins, "birth_date", None) else ""
    return {
        "ref_number":   ins.ref_number or getattr(ins, "ref", "") or "",
        "full_name":    full_name,
        "birth_date":   _fmt_d(ins.birth_date),       # keep if some templates still use it
        "birth_year":   str(birth_year),
        "gender":       ins.gender or "",
        "id_number":    ins.id_number or "",
        "claim_number": ins.claim_number or "",
        "phone": ins.phone or "",
        "address": ins.address or "",
        "city": ins.city or "",
        "age":          _calc_age(ins.birth_date),
        "injury_type":  getattr(ins, "injury_type", "") or "",
    }

def get_report_context(report_id: int, *, insured_id: int | None = None, overrides: dict | None = None) -> dict:
    overrides = overrides or {}
    db_fields  = _fetch_insured_row(insured_id) if insured_id else {}

    # activity date (same as before)
    raw_date   = overrides.get("activity_date", "")
    act_iso    = _to_iso(raw_date) or raw_date

    ctx_fields = {
        "activity_date": act_iso,
        "activity_date_dots": _iso_to_dots(act_iso),
        "surv_place": overrides.get("surv_place", ""),
        "surv_city": overrides.get("surv_city", ""),
        "injury_type": overrides.get("injury_type", ""),
    }

    # === Menora Life Follow-up additional context ===
    ctx_fields.update({
        "injury_type": overrides.get("injury_type", ""),
    })

    # ---- pull Menora Life fields from overrides ----
    bg              = overrides.get("background", "")
    occupation      = overrides.get("occupation", "")
    social_media    = overrides.get("social_media", "")
    social_ident    = overrides.get("social_media_identification", "")
    tracking_date   = overrides.get("tracking_date", "")
    start_time      = overrides.get("start_time", "")
    end_time        = overrides.get("end_time", "")
    summary         = overrides.get("summary", "")
    authorities_1   = overrides.get("authorities_1", "")
    dnb             = overrides.get("dnb", "")          # 👈 NEW
    authorities_2   = overrides.get("authorities_2", "")

    db_fields.update({
        "background": bg,
        "occupation": occupation,
        "social_media": social_media,
        "social_media_identification": social_ident,
        "tracking_date": tracking_date,
        "start_time": start_time,
        "end_time": end_time,
        "summary": summary,
        "authorities_1": authorities_1,
        "dnb": dnb,
        "authorities_2": authorities_2,
    })

    phone_override = overrides.get("phone", "")
    if phone_override:
        db_fields["phone"] = phone_override

    address_override = overrides.get("address", "")
    if address_override:
        db_fields["address"] = address_override

    # 🔹 DEBUG: what’s going into db.background?
    try:
        current_app.logger.info(
            "[CTX] get_report_context: db.background=%r (report_id=%s, insured_id=%s)",
            db_fields.get("background", ""),
            report_id,
            insured_id,
        )
    except Exception:
        pass

    lex = _gender_lex(db_fields.get("gender"))

    return {
        "db": db_fields,
        "ctx": ctx_fields,
        "lex": lex,
        "now": _now_hebrew(),
    }


def _apply_ref_and_version(ctx, ref_number: str = "", reference_no: str = "", version_no=None):
    """
    Normalize 'מספרנו' across all reports.

    Rule:
      • version == 0  ->  base ref only   (e.g. 67590)
      • version > 0   ->  base.version    (e.g. 67590.2)

    base ref is taken from ctx['db']['ref_number'] or the explicit ref_number
    parameter. The final value is written to BOTH:

      ctx['reference_no']
      ctx['db']['ref_number']      # used by {{ db.ref_number }} in templates
    """
    v = int(version_no or 0) if version_no is not None else 0
    db = ctx.setdefault("db", {})

    base = (db.get("ref_number") or ref_number or "").strip()

    if reference_no:
        final = reference_no.strip()
    elif base:
        final = f"{base}.{v}" if v > 0 else base
    else:
        final = ""

    if final:
        ctx["reference_no"] = final
        db["ref_number"] = final

    ctx["version_no"] = v
    return ctx


def _collect_overrides_from_query(args) -> dict:
    """
    Read known inputs from request.args and build a nested overrides dict.
    Empty values are ignored (keep defaults).
    """
    mapping = {
        # ---- generic fields (existing) ----
        "activity_date": "case.activity_date",
        "claim_number": "case.claim_number",
        "insured_name": "insured.name",
        "insured_id": "insured.id",
        "insured_phone": "insured.phone",
        "insured_address": "insured.address",
        "injury_type": "insured.injury_type",
        "surv_place": "surveillance.place",
        "surv_city": "surveillance.city",

        # ---- Menora Life follow-up fields ----
        "background": "db.background",
        "occupation": "insured.occupation",
        "social_media": "insured.social_media",
        "social_media_identification": "insured.social_media_identification",
        "tracking_date": "case.tracking_date",
        "start_time": "case.start_time",
        "end_time": "case.end_time",
        "summary": "case.summary",
        "authorities_1": "case.authorities_1",
        "authorities_2": "case.authorities_2",
        "dnb": "db.dnb",
    }

    out = {}
    for qkey, path in mapping.items():
        val = (args.get(qkey) or "").strip()
        if val:
            _deep_set(out, path, val)
    return out


def _collect_overrides_from_json(json_body: dict | None) -> dict:
    """Same as query, but from a JSON body for the download route."""
    if not json_body:
        return {}
    fields = {
        "activity_date": "case.activity_date",
        "claim_number": "case.claim_number",
        "insured_name": "insured.name",
        "insured_id": "insured.id",
        "insured_phone": "insured.phone",
        "insured_address": "insured.address",
        "injury_type": "insured.injury_type",
        "surv_place": "surveillance.place",
        "surv_city": "surveillance.city",

        # --- Menora Life Follow-up extra fields ---
        "background": "db.background",
        "occupation": "insured.occupation",
        "social_media": "insured.social_media",
        "social_media_identification": "insured.social_media_identification",
        "tracking_date": "case.tracking_date",
        "start_time": "case.start_time",
        "end_time": "case.end_time",
        "summary": "case.summary",
    }

    out = {}
    for k, dotted in fields.items():
        val = (json_body.get(k) or "").strip()
        if val:
            _deep_set(out, dotted, val)
    return out


def ensure_output_dir():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    # helpful once at boot to confirm where files go
    try:
        current_app.logger.info(f"[DOCX] OUTPUT_DIR={OUTPUT_DIR}")
    except Exception:
        pass

def load_template_docx(template_name: str) -> str:
    """
    Returns absolute path to the DOCX template, using Config.DOCX_TEMPLATES_DIR.
    """
    path = os.path.join(DOCX_TEMPLATES_DIR, template_name)
    if not os.path.exists(path):
        raise FileNotFoundError(f"Template not found: {path}")
    # optional: log it for sanity
    try:
        current_app.logger.info(f"[DOCX] Using template: {path}")
    except Exception:
        pass
    return path


def render_docx_bytes(template_path: str, context: dict) -> bytes:
    """
    Renders a DOCX from template + context and returns it as bytes.
    Expects the DOCX to use real Jinja keys (e.g., {{ case.report_date }}).
    """
    tpl = DocxTemplate(template_path)
    tpl.render(context)
    buf = io.BytesIO()
    tpl.save(buf)
    return buf.getvalue()


def build_resized_inline_image(
    doc,
    image_path,
    *,
    max_width_mm: float,
    max_height_mm: float
):
    """
    Create a resized InlineImage that fits within max width/height (mm),
    preserving aspect ratio. Matches existing Siudi behavior.
    """
    try:
        with Image.open(image_path) as im:
            w_px, h_px = im.size
    except Exception:
        # Fallback: let docx handle it (should not happen)
        return InlineImage(doc, image_path)

    # portrait vs landscape
    if h_px >= w_px:
        # portrait → constrain height
        return InlineImage(
            doc,
            image_path,
            height=Mm(max_height_mm)
        )
    else:
        # landscape → constrain width
        return InlineImage(
            doc,
            image_path,
            width=Mm(max_width_mm)
        )

def prepare_photos(
    doc,
    image_paths,
    *,
    max_width_mm: float,
    max_height_mm: float
):
    """
    Load + resize photos and return PreparedPhoto objects.
    No layout logic here.
    """
    prepared = []

    for path in image_paths:
        try:
            with Image.open(path) as im:
                w_px, h_px = im.size
                orientation = "portrait" if h_px >= w_px else "landscape"
        except Exception:
            orientation = "landscape"

        opt_path = optimize_image_for_docx(path)

        img = build_resized_inline_image(
            doc,
            opt_path,
            max_width_mm=max_width_mm,
            max_height_mm=max_height_mm
        )

        prepared.append(
            PreparedPhoto(
                path=path,
                orientation=orientation,
                image=img
            )
        )

    return prepared



def build_photo_pages(prepared_photos):
    """
    Build page-based photo layout.
    Each page contains 1 or 2 photos, in strict chronological order.

    Layout rules:
    - portrait + portrait  → side-by-side (row)
    - any other combo      → stacked (column)
    """

    pages = []
    i = 0

    while i < len(prepared_photos):
        first = prepared_photos[i]
        second = prepared_photos[i + 1] if i + 1 < len(prepared_photos) else None

        # Single photo page
        if not second:
            pages.append({
                "layout": "single",
                "photos": [first.image],
                "is_two_portrait": False,
            })
            break

        is_same_orientation = (
                first.orientation == second.orientation
        )

        pages.append({
            "layout": "row" if is_same_orientation else "column",
            "photos": [first.image, second.image],
            "is_two_portrait": (
                    first.orientation == "portrait"
                    and second.orientation == "portrait"
            ),
        })

        i += 2

    return pages


def optimize_image_for_docx(
    src_path: str,
    max_px: int = 1600,
    quality: int = 75
) -> str:
    """
    Downscale + recompress image for DOCX embedding.
    Returns path to optimized temp JPEG.
    """
    img = Image.open(src_path)

    # Convert to RGB (important for PNG / WEBP)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")

    # Resize (preserve aspect ratio)
    img.thumbnail((max_px, max_px), Image.LANCZOS)

    # Save to temp JPEG
    fd, tmp_path = tempfile.mkstemp(suffix=".jpg")
    os.close(fd)

    img.save(
        tmp_path,
        format="JPEG",
        quality=quality,
        optimize=True,
        progressive=True
    )

    return tmp_path


############################### ROUTES #############################################
@reports_docx_bp.route("/<int:report_id>/download/<path:filename>", methods=["GET"])
def download_docx(report_id: int, filename: str):
    """
    Download a previously generated DOCX.
    """
    abs_path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(abs_path):
        abort(404)
    return send_file(abs_path, as_attachment=True, download_name=filename, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# --- add near other imports ---
import subprocess, tempfile, shutil
from flask import send_file

def _resolve_soffice_path() -> str:
    """
    Prefer soffice.exe on Windows; fall back to PATH.
    """
    from .config import Config
    cand = getattr(Config, "LIBREOFFICE_BIN", None)
    if cand and os.path.exists(cand):
        # normalize to .exe if someone pointed at soffice.com
        if cand.lower().endswith("soffice.com"):
            exe = cand[:-4] + "exe"
            if os.path.exists(exe):
                return exe
        return cand

    # Common Windows installs
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        os.path.expandvars(r"%LOCALAPPDATA%\Programs\LibreOffice\program\soffice.exe"),
    ]
    for c in candidates:
        if os.path.exists(c):
            return c

    # PATH (may resolve to soffice.exe already)
    from shutil import which
    w = which("soffice")
    if w:
        # if it’s .com and .exe exists next to it, prefer .exe
        if w.lower().endswith("soffice.com"):
            exe = w[:-4] + "exe"
            if os.path.exists(exe):
                return exe
        return w

    return "soffice"


def _docx_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """
    Robust DOCX -> PDF on Windows/LibreOffice:
    - uses a temp user profile (avoids locks)
    - runs in the temp dir as CWD
    - tries 'pdf' then 'pdf:writer_pdf_Export'
    - scans for any *.pdf (LO sometimes renames)
    - waits for the file to actually appear
    """
    import subprocess, tempfile, shutil, time, glob
    from pathlib import Path

    soffice = _resolve_soffice_path()

    tmpdir = tempfile.mkdtemp(prefix="lo_")
    tmp = Path(tmpdir)
    try:
        docx_path = tmp / "report.docx"
        expected_pdf = tmp / "report.pdf"
        docx_path.write_bytes(docx_bytes)

        profile_url = tmp.as_uri() + "/lo_profile"

        def try_convert(filter_spec: str) -> tuple[int, str, str]:
            cmd = [
                soffice,
                "--headless", "--nologo", "--nodefault", "--nolockcheck",
                "--norestore", "--invisible",
                f"-env:UserInstallation={profile_url}",
                "--convert-to", filter_spec,
                "--outdir", str(tmp),
                str(docx_path),
            ]

            env = os.environ.copy()
            env["SAL_DISABLE_OPENCL"] = "true"
            env["LO_JAVA_ENABLED"] = "false"

            proc = subprocess.run(
                cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                text=True, cwd=str(tmp), env=env
            )
            # wait up to 15s for *any* PDF to appear
            deadline = time.time() + 15.0
            while time.time() < deadline:
                if expected_pdf.exists():
                    break
                pdfs = list(tmp.glob("*.pdf"))
                if pdfs:
                    break
                time.sleep(0.2)
            return proc.returncode, proc.stdout, proc.stderr

        rc, out, err = try_convert("pdf")
        pdfs = sorted(tmp.glob("*.pdf"), key=lambda p: p.stat().st_mtime)
        if not pdfs:
            # second attempt with explicit filter
            rc2, out2, err2 = try_convert("pdf:writer_pdf_Export")
            out += "\n\n-- second attempt stdout --\n" + out2
            err += "\n\n-- second attempt stderr --\n" + err2
            pdfs = sorted(tmp.glob("*.pdf"), key=lambda p: p.stat().st_mtime)

        if not pdfs:
            listing = ""
            try:
                listing = "\n".join(p.name for p in tmp.iterdir())
            except Exception:
                listing = "(dir listing failed)"
            raise RuntimeError(
                "LibreOffice convert returned rc="
                f"{rc}, but no PDF found.\n"
                f"STDOUT:\n{out}\n\nSTDERR:\n{err}\n\n"
                f"TEMP DIR: {tmp}\nFILES:\n{listing}\n"
            )

        pdf_path = pdfs[-1]
        time.sleep(0.2)  # flush grace on Windows
        return pdf_path.read_bytes()

    finally:
        # try multiple times in case Windows locks the temp files briefly
        for _ in range(6):
            try:
                shutil.rmtree(tmpdir)
                break
            except Exception:
                time.sleep(0.3)


# ---- preview PDF (REPLACE with photo-aware) ----
# ====== imports this block relies on (ensure these exist at top of file) ======
# import os, io, tempfile
# from flask import request, send_file, current_app
# from docxtpl import DocxTemplate, InlineImage
# from docx.shared import Mm
# from PIL import Image, ImageOps   # if Pillow not installed, pip install pillow

# ====== small helpers (drop in once) =========================================
# ---- helpers: tmpdir, orientation, resize, preprocess (with logging) ----
from contextlib import contextmanager
from PIL import Image, ImageOps  # make sure Pillow is installed

@contextmanager
def _tmpdir(prefix="report_img_"):
    td = tempfile.TemporaryDirectory(prefix=prefix)
    try:
        yield td.name
    finally:
        td.cleanup()

def _detect_orientation(path: str) -> str:
    """Return 'portrait' if height >= width*1.10 after EXIF rotate, else 'landscape'."""
    try:
        with Image.open(path) as im:
            im = ImageOps.exif_transpose(im)
            w, h = im.size
            return "portrait" if (h / max(1, w)) >= 1.10 else "landscape"
    except Exception:
        return "landscape"

def _resize_copy(src: str, dst: str, max_w: int, max_h: int, quality: int = 82) -> tuple[int,int,int,int]:
    """
    Resize into dst (JPEG) with EXIF rotation, keep aspect.
    Returns (orig_w, orig_h, new_w, new_h).
    """
    with Image.open(src) as im:
        im = ImageOps.exif_transpose(im)
        ow, oh = im.size
        im.thumbnail((max_w, max_h), Image.LANCZOS)
        nw, nh = im.size
        if im.mode not in ("RGB", "L"):
            im = im.convert("RGB")
        im.save(dst, format="JPEG",
                quality=quality, optimize=True, progressive=True, subsampling="4:2:0")
        return ow, oh, nw, nh

def _preprocess_for_docx(tmpdir: str, paths: list[str]) -> list[dict]:
    """
    For each path, produce a dict:
      { "path": <processed_path>, "guess": "portrait"/"landscape",
        "ow": <orig_w>, "oh": <orig_h>, "w": <new_w>, "h": <new_h> }
    """
    out: list[dict] = []
    for idx, p in enumerate(paths):
        guess = _detect_orientation(p)
        # caps: keep files small but clear
        max_w, max_h = ((700, 1100) if guess == "portrait" else (1400, 900))
        dst = os.path.join(tmpdir, f"img_{idx:03d}.jpg")
        try:
            ow, oh, nw, nh = _resize_copy(p, dst, max_w, max_h, quality=80)
            current_app.logger.info(
                "[PREPROC] %s -> %s | guess=%s | orig=%dx%d -> proc=%dx%d",
                os.path.basename(p), os.path.basename(dst), guess, ow, oh, nw, nh
            )
            out.append({"path": dst, "guess": guess, "ow": ow, "oh": oh, "w": nw, "h": nh})
        except Exception as e:
            current_app.logger.warning("[PREPROC] failed %s: %s; using original", p, e)
            # fall back to original (we don’t know processed size, so re-open to get it)
            try:
                with Image.open(p) as im:
                    im = ImageOps.exif_transpose(im)
                    nw, nh = im.size
            except Exception:
                nw = nh = 0
            out.append({"path": p, "guess": guess, "ow": 0, "oh": 0, "w": nw, "h": nh})
    return out



def _report_media_root() -> str:
    return current_app.config.get(
        "REPORT_MEDIA_DIR",
        os.path.join(current_app.instance_path, "report_media")
    )

# ====== FULL ROUTE: /reports/<id>/preview-pdf  ===============================

# ====== FULL ROUTE: /reports/<id>/preview-pdf  ===============================

@reports_docx_bp.route("/<int:report_id>/preview-pdf", methods=["GET"])
def preview_docx_as_pdf(report_id: int):
    import io, os
    from PIL import Image
    from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Mm

    insured_id = request.args.get("insured_id", type=int)

    # NEW: reference + version
    ref_number   = (request.args.get("ref_number") or "").strip()
    reference_no = (request.args.get("reference_no") or "").strip()
    version_no   = request.args.get("version_no", type=int)

    # ---- collect overrides ----
    overrides = {
        "activity_date": (request.args.get("activity_date") or "").strip(),
        "surv_place":    (request.args.get("surv_place") or "").strip(),
        "surv_city":     (request.args.get("surv_city")  or "").strip(),
        "injury_type":   (request.args.get("injury_type") or "").strip(),
        "background":    (request.args.get("background") or "").strip(),
        "occupation":    (request.args.get("occupation") or "").strip(),
        "social_media":  (request.args.get("social_media") or "").strip(),
        "social_media_identification":
            (request.args.get("social_media_identification") or "").strip(),
        "tracking_date": (request.args.get("tracking_date") or "").strip(),
        "start_time":    (request.args.get("start_time") or "").strip(),
        "end_time":      (request.args.get("end_time") or "").strip(),
        "summary":       (request.args.get("summary") or "").strip(),
        "authorities_1": (request.args.get("authorities_1") or "").strip(),
        "authorities_2": (request.args.get("authorities_2") or "").strip(),
        "phone":         (request.args.get("phone") or "").strip(),
        "address": (request.args.get("address") or "").strip(),
        "dnb":           (request.args.get("dnb") or "").strip(),
    }

    ctx = get_report_context(report_id, insured_id=insured_id, overrides=overrides)

    ctx = _apply_ref_and_version(
        ctx,
        ref_number=ref_number,
        reference_no=reference_no,
        version_no=version_no,
    )

    # ------------------------------------------------------------
    # Tracking date formatted for Menora header: dd-mm-yyyy
    # ------------------------------------------------------------
    db = ctx.get("db", {})
    ctx["tracking_date_ddmmyyyy"] = _iso_to_ddmmyyyy_dash(
        db.get("tracking_date", "")
    )

    # ------------------------------------------------------------
    # 🔒 RESET PHOTO CONTEXT (prevents leakage between renders)
    # ------------------------------------------------------------
    ctx["photo_pages"] = []
    ctx["social_photos"] = []
    ctx["authorities_table_photo"] = ""
    ctx["authorities_table_photo_2"] = ""

    # ------------------------------------------------------------
    #          TEMPLATE DETECTION
    # ------------------------------------------------------------
    tmpl_key = (request.args.get("template", "siudi") or "siudi").strip().lower()
    template_path = load_template_docx(map_template_key(tmpl_key))
    current_app.logger.info("[DEBUG][PHOTOS] Using template key=%s -> %s", tmpl_key, template_path)

    tpl = DocxTemplate(template_path)

    # ------------------------------------------------------------
    #      SIUDI INVOICE — populate invoice fields (PREVIEW)
    # ------------------------------------------------------------
    if tmpl_key == "siudi_invoice":
        # Ensure required objects exist (🔥 prevents UndefinedError)
        ctx.setdefault("claim", {})
        ctx.setdefault("insured", {})
        ctx.setdefault("totals", {})

        # ---- Invoice header ----
        inv_date_str = (request.args.get("inv_date") or "").strip()
        inv_iso = _to_iso(inv_date_str) or inv_date_str
        ctx["inv_date"] = ddmmyyyy(inv_iso) if "-" in (inv_iso or "") else inv_date_str

        ctx["inv_number"] = (request.args.get("inv_number") or "").strip()
        ctx["inv_ref"] = (request.args.get("inv_ref") or "").strip()

        # ---- Claim ----
        ctx["claim"]["number"] = (
                request.args.get("claim.number")
                or ctx.get("db", {}).get("claim_number", "")
        )
        ctx["claim"]["subject"] = (
                request.args.get("claim.subject")
                or ctx.get("db", {}).get("full_name", "")
        )

        # ---- Insured ----
        ctx["insured"]["id_number"] = (
                request.args.get("insured.id_number")
                or ctx.get("db", {}).get("id_number", "")
        )

        # ---- Totals ----
        ctx["totals"]["subtotal"] = (request.args.get("totals_subtotal") or "").strip()
        ctx["totals"]["vat_rate"] = (request.args.get("totals_vat_rate") or "").strip()
        ctx["totals"]["vat_amount"] = (request.args.get("totals_vat_amount") or "").strip()
        ctx["totals"]["total"] = (request.args.get("totals_total") or "").strip()

    # ------------------------------------------------------------
    #            FOLLOW-UP specific fields
    # ------------------------------------------------------------
    if tmpl_key == "menora_life_followup":
        raw = (request.args.get("tracking_raw") or "").strip()
        ctx["tracking_rows"] = _parse_tracking_rows(raw) if raw else []

    if tmpl_key == "menora_life_photos":
        iso = (request.args.get("photo_date") or "").strip()
        if iso:
            ctx.setdefault("db", {})["photo_date"] = ddmmyyyy(iso)

    # ------------------------------------------------------------
    #              RESOLVE PHOTO FILENAMES
    # ------------------------------------------------------------
    # --- common media root + resolver (photos + 'טבלת רשויות') ---

    case_id = str(insured_id or "").strip()
    rep_id  = str(report_id)

    media_root = current_app.config.get(
        "REPORT_MEDIA_DIR",
        os.path.join(current_app.instance_path, "report_media"),
    )

    # The “normal” place where local-dropbox photos are stored
    base_dir = os.path.join(media_root, case_id, rep_id)

    # Where to search first (local storage)
    candidate_dirs = [
        base_dir,                       # <media_root>/<insured>/<report>
        os.path.join(media_root, case_id)  # <media_root>/<insured> (legacy / safety)
    ]

    def resolve_one(name: str) -> str | None:
        """
        Resolve a selected image name to a local filesystem path.

        Supports:
          1) already-local files under REPORT_MEDIA_DIR/<insured>/<report>
          2) Dropbox-backed GilMedia rows (downloads to a local cache on-demand)
        """
        if not name:
            return None

        base = os.path.basename(name)

        # 1) Already stored locally (old/local-dropbox flow)
        for d in candidate_dirs:
            p = os.path.join(d, base)
            if os.path.exists(p):
                return p

        # 2) Dropbox-backed media: download to a cache folder and return that path
        try:
            insured_int = int(case_id)
        except Exception:
            insured_int = None

        if not insured_int:
            return None

        # ✅ IMPORTANT: db is not guaranteed to exist in this module unless imported
        try:
            from . import db  # <-- your Flask SQLAlchemy instance
        except Exception as ex:
            current_app.logger.warning("[PHOTOS] cannot import db in reports_docx.py: %s", ex)
            return None

        # Fetch GilMedia row by insured_id + file_name
        media_row = None
        try:
            from .models import GilMedia
            media_row = (
                db.session.query(GilMedia)
                .filter(
                    GilMedia.insured_id == insured_int,
                    GilMedia.file_name == base
                )
                .order_by(GilMedia.media_id.desc())
                .first()
            )
        except Exception as ex:
            current_app.logger.warning("[PHOTOS] GilMedia query failed for %s: %s", base, ex)
            media_row = None

        if not media_row:
            current_app.logger.info("[PHOTOS] no GilMedia row for file_name=%s insured_id=%s", base, insured_int)
            return None

        dropbox_path = getattr(media_row, "dropbox_path", None)
        if not dropbox_path:
            current_app.logger.info("[PHOTOS] GilMedia row found but missing dropbox_path for %s", base)
            return None

        # Local cache path
        cache_dir = os.path.join(base_dir, "_dropbox_cache")
        os.makedirs(cache_dir, exist_ok=True)
        cache_path = os.path.join(cache_dir, base)

        if os.path.exists(cache_path):
            return cache_path

        # Download from Dropbox using the same client as routes.py
        try:
            try:
                from .routes import dbx  # lazy import
            except Exception as ex:
                current_app.logger.warning("[PHOTOS] cannot import dbx from routes.py: %s", ex)
                return None

            current_app.logger.info("[PHOTOS] downloading from Dropbox: %s -> %s", dropbox_path, cache_path)
            dbx.files_download_to_file(cache_path, dropbox_path)

            if os.path.exists(cache_path):
                return cache_path

            current_app.logger.warning("[PHOTOS] download finished but file missing: %s", cache_path)
            return None

        except Exception as ex:
            current_app.logger.warning("[PHOTOS] Dropbox download failed for %s: %s", base, ex)
            return None


    # ------------------------------------------------------------
    #            DEBUG: which photo sets received?
    # ------------------------------------------------------------
    photos_generic = request.args.getlist("selected_photos[]")
    photos_life    = request.args.getlist("selected_life_photos[]")

    current_app.logger.info("[DEBUG][PHOTOS] tmpl_key=%s", tmpl_key)
    current_app.logger.info("[DEBUG][PHOTOS] selected_photos[]       = %s", photos_generic)
    current_app.logger.info("[DEBUG][PHOTOS] selected_life_photos[] = %s", photos_life)

    # ------------------------------------------------------------
    #            Decide which photo list to use
    # ------------------------------------------------------------
    if tmpl_key == "menora_life_followup":
        active_list = photos_life
    elif tmpl_key == "menora_life_photos":
        active_list = photos_generic
    else:
        active_list = photos_generic

    # ------------------------------------------------------------
    #            Resolve paths + debug log each
    # ------------------------------------------------------------
    resolved_paths = []
    for fn in active_list:
        path = resolve_one(fn)
        current_app.logger.info("[DEBUG][PHOTOS] resolve_one('%s') -> %s", fn, path)
        if not path:
            current_app.logger.warning("[PHOTOS] not found: %s", fn)
        else:
            resolved_paths.append(path)

    # ------------------------------------------------------------
    #      FOLLOW-UP: flat list goes to {{ social_photos }}
    # ------------------------------------------------------------
    if tmpl_key == "menora_life_followup":
        prepared = prepare_photos(
            tpl,
            resolved_paths,
            max_width_mm=155,
            max_height_mm=145
        )

        photo_pages = build_photo_pages(prepared)

        # TEMP: expose to template
        ctx["photo_pages"] = photo_pages

        # For now: flat list, same as before
        ctx["social_photos"] = [p.image for p in prepared]

        # Authorities tables
        for key, placeholder in [
            ("authorities_table", "authorities_table_photo"),
            ("authorities_table_2", "authorities_table_photo_2"),
        ]:
            name = (request.args.get(key) or "").strip()
            img_path = resolve_one(name)
            # ctx[placeholder] = InlineImage(tpl, img_path, width=Mm(155)) if img_path else ""
            if img_path:
                opt_path = optimize_image_for_docx(img_path)
                ctx[placeholder] = InlineImage(tpl, opt_path, width=Mm(155))
            else:
                ctx[placeholder] = ""

    # ------------------------------------------------------------
    # SIUDI + GENERIC PHOTO REPORTS — unified handling
    # ------------------------------------------------------------
    if tmpl_key != "menora_life_followup":
        prepared = prepare_photos(
            tpl,
            resolved_paths,
            max_width_mm=155,
            max_height_mm=130
        )

        photo_pages = build_photo_pages(prepared)
        ctx["photo_pages"] = photo_pages
        ctx["social_photos"] = [p.image for p in prepared]


    # ------------------------------------------------------------
    #     MENORA LIFE INVOICE — populate invoice fields (PREVIEW)
    # ------------------------------------------------------------
    if tmpl_key == "menora_life_invoice":
        ctx.setdefault("claim", {})
        ctx.setdefault("insured", {})

        # Parse invoice date
        inv_date_str = (request.args.get("inv_date") or "").strip()
        inv_iso = _to_iso(inv_date_str) or inv_date_str
        ctx["inv_date"] = ddmmyyyy(inv_iso) if "-" in (inv_iso or "") else inv_date_str

        # Simple scalar fields
        ctx["inv_number"] = (request.args.get("inv_number") or "").strip()
        ctx["inv_ref"] = (request.args.get("inv_ref") or "").strip()
        ctx["life_followup_date"] = (request.args.get("life_followup_date") or "").strip()

        # Claim / insured (needed for header)
        ctx["claim"]["number"] = request.args.get("claim.number", "") or ctx["db"].get("claim_number", "")
        ctx["claim"]["subject"] = request.args.get("claim.subject", "") or ctx["db"].get("full_name", "")
        ctx["insured"]["id_number"] = request.args.get("insured.id_number", "") or ctx["db"].get("id_number", "")

        # Totals
        ctx["life_subtotal"] = (request.args.get("life_subtotal") or "").strip()
        ctx["life_vat_amount"] = (request.args.get("life_vat_amount") or "").strip()
        ctx["life_total"] = (request.args.get("life_total") or "").strip()

        # Items (life_items[1][text] etc)
        life_items = []
        for i in range(1, 50):
            text = request.args.get(f"life_items[{i}][text]")
            amount = request.args.get(f"life_items[{i}][amount]")
            if text or amount:
                life_items.append({
                    "text": text or "",
                    "amount": amount or "",
                })
        ctx["life_items"] = life_items

    # ------------------------------------------------------------
    # GENERATE PDF
    # ------------------------------------------------------------
    buf = io.BytesIO()

    tpl.render(ctx)
    tpl.save(buf)
    pdf_bytes = _docx_to_pdf_bytes(buf.getvalue())
    return send_file(io.BytesIO(pdf_bytes),
                     mimetype="application/pdf",
                     download_name="preview.pdf")



@reports_docx_bp.route("/<int:report_id>/render-docx", methods=["POST"])
def render_docx_download(report_id: int):
    import io, os, datetime, json
    from PIL import Image
    from flask import current_app, request, jsonify, url_for
    from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Mm

    payload = request.get_json(silent=True) or {}

    insured_id = payload.get("insured_id")

    # NEW: reference + version from client
    ref_number   = (payload.get("ref_number") or "").strip()
    reference_no = (payload.get("reference_no") or "").strip()
    version_no   = payload.get("version_no")

    # ---- collect overrides from JSON (including background / dnb / phone) ----
    overrides = {
        "activity_date": payload.get("activity_date", "").strip(),
        "surv_place":    payload.get("surv_place", "").strip(),
        "surv_city":     payload.get("surv_city",  "").strip(),
        "injury_type":   payload.get("injury_type", "").strip(),

        # Menora Life follow-up fields
        "background":    payload.get("background", "").strip(),
        "occupation":    payload.get("occupation", "").strip(),
        "social_media":  payload.get("social_media", "").strip(),
        "social_media_identification":
            payload.get("social_media_identification", "").strip(),
        "tracking_date": payload.get("tracking_date", "").strip(),
        "start_time":    payload.get("start_time", "").strip(),
        "end_time":      payload.get("end_time", "").strip(),
        "summary":       payload.get("summary", "").strip(),
        "authorities_1": payload.get("authorities_1", "").strip(),
        "authorities_2": payload.get("authorities_2", "").strip(),
        "phone":         payload.get("phone", "").strip(),
        "address": payload.get("address", "").strip(),
        "dnb":           payload.get("dnb", "").strip(),
    }

    try:
        current_app.logger.info(
            "[CTX][download] payload.background=%r (report_id=%s, insured_id=%s)",
            overrides["background"],
            report_id,
            insured_id,
        )
    except Exception:
        pass

    # base context (db + ctx + lex + now)
    ctx = get_report_context(report_id, insured_id=insured_id, overrides=overrides)

    # normalise ref/version into ctx["db"]["ref_number"]
    ctx = _apply_ref_and_version(
        ctx,
        ref_number=ref_number,
        reference_no=reference_no,
        version_no=version_no,
    )

    # ------------------------------------------------------------
    # Tracking date formatted for Menora header: dd-mm-yyyy
    # ------------------------------------------------------------
    db = ctx.get("db", {})
    ctx["tracking_date_ddmmyyyy"] = _iso_to_ddmmyyyy_dash(
        db.get("tracking_date", "")
    )

    # ------------------------------------------------------------
    # 🔒 RESET PHOTO CONTEXT (prevents leakage between renders)
    # ------------------------------------------------------------
    ctx["photo_pages"] = []
    ctx["social_photos"] = []
    ctx["authorities_table_photo"] = ""
    ctx["authorities_table_photo_2"] = ""


    # --- template path ---
    tmpl_key = (payload.get("template") or "siudi").strip().lower()
    template_path = load_template_docx(map_template_key(tmpl_key))
    current_app.logger.info("[DOCX] Using template: %s (key=%s)", template_path, tmpl_key)

    tpl = DocxTemplate(template_path)

    # --- Menora life follow-up: tracking activities table ---
    if tmpl_key == "menora_life_followup":
        raw = (payload.get("tracking_raw") or "").strip()
        ctx["tracking_rows"] = _parse_tracking_rows(raw) if raw else []

    # Menora Life – Photos report: pass {{ db.photo_date }} as dd/mm/yyyy
    if tmpl_key == "menora_life_photos":
        iso = (payload.get("photo_date") or "").strip()
        if iso:
            ctx.setdefault("db", {})["photo_date"] = ddmmyyyy(iso)

    # --- common media root + resolver (photos + 'טבלת רשויות') ---
    case_id = str(insured_id or payload.get("case_id", "")).strip()
    rep_id = str(report_id)

    media_root = current_app.config.get(
        "REPORT_MEDIA_DIR",
        os.path.join(current_app.instance_path, "report_media"),
    )

    # The “normal” place where local-dropbox photos are stored
    base_dir = os.path.join(media_root, case_id, rep_id)

    # Where to search first (local storage)
    candidate_dirs = [
        base_dir,  # <media_root>/<insured>/<report>
        os.path.join(media_root, case_id)  # <media_root>/<insured>  (legacy / safety)
    ]

    def resolve_one(name: str) -> str | None:
        """Resolve a selected image name to a local filesystem path.

        Supports:
          1) already-local files under REPORT_MEDIA_ROOT/<insured>/<report>
          2) Dropbox-backed GilMedia rows (downloads to a local cache on-demand)
        """
        if not name:
            return None

        # Normalize (UI sometimes sends full paths)
        base = os.path.basename(name)

        # 1) Already stored locally (old/local-dropbox flow)
        for d in candidate_dirs:
            p = os.path.join(d, base)
            if os.path.exists(p):
                return p

        # 2) Dropbox-backed media: download to a cache folder and return that path
        try:
            insured_int = int(case_id)
        except Exception:
            insured_int = None

        if insured_int:
            media_row = None
            try:
                from .models import GilMedia
                media_row = (
                    db.session.query(GilMedia)
                    .filter(
                        GilMedia.insured_id == insured_int,
                        GilMedia.file_name == base
                    )
                    .order_by(GilMedia.media_id.desc())
                    .first()
                )
            except Exception:
                media_row = None

            if media_row and getattr(media_row, "dropbox_path", None):
                cache_dir = os.path.join(base_dir, "_dropbox_cache")
                os.makedirs(cache_dir, exist_ok=True)
                cache_path = os.path.join(cache_dir, base)

                # already downloaded
                if os.path.exists(cache_path):
                    return cache_path

                # Download from Dropbox
                try:
                    try:
                        from .routes import dbx  # lazy import
                    except Exception:
                        dbx = None

                    if dbx:
                        dbx.files_download_to_file(cache_path, media_row.dropbox_path)
                        if os.path.exists(cache_path):
                            return cache_path
                except Exception:
                    return None

        return None

    # --- Menora Life: two authorities-table images ---
    if tmpl_key == "menora_life_followup":
        def _add_table_from_payload(key: str, placeholder: str):
            tbl_name = (payload.get(key) or "").strip()
            img_path = resolve_one(tbl_name) if tbl_name else None
            if img_path:
                opt_path = optimize_image_for_docx(img_path)
                ctx[placeholder] = InlineImage(tpl, opt_path, width=Mm(155))
            else:
                ctx[placeholder] = ""

        _add_table_from_payload("authorities_table", "authorities_table_photo")
        _add_table_from_payload("authorities_table_2", "authorities_table_photo_2")

    # ===================== SIUDI INVOICE (no photos) ======================
    if tmpl_key == "siudi_invoice":
        inv_date_iso = (payload.get("inv_date") or "").strip()
        ctx.setdefault("insured", {})
        ctx.setdefault("claim", {})
        ctx.setdefault("totals", {})

        ctx.update({
            "inv_date":   ddmmyyyy(inv_date_iso),
            "inv_number": (payload.get("inv_number") or "").strip(),
            "inv_ref":    (payload.get("inv_ref") or "").strip(),
        })
        ctx["insured"]["id_number"] = (payload.get("insured", {}).get("id_number") or "").strip()
        ctx["claim"]["number"]      = (payload.get("claim",   {}).get("number")    or "").strip()
        ctx["claim"]["subject"]     = (payload.get("claim",   {}).get("subject")   or "").strip()
        ctx["totals"].update({
            "subtotal":   (payload.get("totals", {}).get("subtotal")   or "").strip(),
            "vat_rate":   (payload.get("totals", {}).get("vat_rate")   or "").strip(),
            "vat_amount": (payload.get("totals", {}).get("vat_amount") or "").strip(),
            "total":      (payload.get("totals", {}).get("total")      or "").strip(),
        })

        out_dir = os.path.join(current_app.instance_path, "generated_reports")
        os.makedirs(out_dir, exist_ok=True)

        stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"report_{report_id}_{stamp}.docx"
        abs_path = os.path.join(out_dir, filename)

        tpl.render(ctx)
        tpl.save(abs_path)

        docx_url = url_for("generated_reports", filename=filename)
        return jsonify({"ok": True, "docx_url": docx_url})
    # =================== /SIUDI INVOICE ===================================

    # ===================== MENORA LIFE INVOICE ============================
    if tmpl_key == "menora_life_invoice":
        inv_date_str = (payload.get("inv_date") or "").strip()
        followup_date_text = (payload.get("life_followup_date") or "").strip()

        inv_iso = _to_iso(inv_date_str) or inv_date_str
        inv_date_dmy = ddmmyyyy(inv_iso) if "-" in (inv_iso or "") else inv_date_str

        life_items = payload.get("life_items") or []
        if isinstance(life_items, str):
            try:
                life_items = json.loads(life_items)
            except Exception:
                life_items = []

        totals = payload.get("totals") or {}

        ctx.setdefault("insured", {})
        ctx.setdefault("claim", {})

        ctx.update({
            "inv_date": inv_date_dmy,
            "inv_number": (payload.get("inv_number") or "").strip(),
            "inv_ref": (payload.get("inv_ref") or "").strip(),
            "life_followup_date": followup_date_text,
            "life_items": life_items,
            "life_subtotal": (totals.get("subtotal") or "").strip(),
            "life_vat_amount": (totals.get("vat_amount") or "").strip(),
            "life_total": (totals.get("total") or "").strip(),
        })

        # 🔥 FIX: ensure claim + insured objects exist
        ctx["claim"]["number"] = payload.get("claim", {}).get("number") or ctx["db"].get("claim_number", "")
        ctx["claim"]["subject"] = payload.get("claim", {}).get("subject") or ctx["db"].get("full_name", "")
        ctx["insured"]["id_number"] = payload.get("insured", {}).get("id_number") or ctx["db"].get("id_number", "")

        # Save DOCX
        out_dir = os.path.join(current_app.instance_path, "generated_reports")
        os.makedirs(out_dir, exist_ok=True)

        stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"report_{report_id}_{stamp}.docx"
        abs_path = os.path.join(out_dir, filename)

        tpl.render(ctx)
        tpl.save(abs_path)

        docx_url = url_for("generated_reports", filename=filename)
        return jsonify({"ok": True, "docx_url": docx_url})

    # =================== /MENORA LIFE INVOICE =============================

    # --- collect selected names from payload (for photo-based reports) ---
    def _normalize_photo_name(v: str) -> str | None:
        if not v:
            return None

        # URL case: /reports/photos/serve?...&name=FILE.png
        if "name=" in v:
            return v.split("name=", 1)[1].split("&", 1)[0]

        # Strip querystring if exists
        if "?" in v:
            v = v.split("?", 1)[0]

        return os.path.basename(v) or None

    raw_names = payload.get("selected_photos") or []

    # fallback for older payloads
    if not raw_names and isinstance(payload.get("photos"), list):
        raw_names = [p.get("name") for p in payload["photos"] if p.get("name")]

    names = [_normalize_photo_name(v) for v in raw_names]
    names = [n for n in names if n]

    paths = [p for n in names if (p := resolve_one(n))]

    # disable fallback for Menora Life follow-up
    if not paths and tmpl_key not in {"menora_life_followup"}:
        folder = os.path.join(media_root, case_id, rep_id)
        if os.path.isdir(folder):
            for fn in sorted(os.listdir(folder)):
                if fn.lower().endswith((".jpg", ".jpeg", ".png", ".webp", ".bmp", ".gif")):
                    paths.append(os.path.join(folder, fn))

    # --- Social media photos for Menora Life follow-up (download) ---
    if tmpl_key == "menora_life_followup":
        # life-followup uses ONLY selected_life_photos[] for the social media block
        ctx["social_photos"] = [
            InlineImage(tpl, optimize_image_for_docx(p), width=Mm(155))
            for p in paths
        ]

        # --- Authorities 1 + 2 (must not depend on active_list!) ---
        for key, placeholder in [
            ("authorities_table", "authorities_table_photo"),
            ("authorities_table_2", "authorities_table_photo_2"),
        ]:
            name = (request.args.get(key) or "").strip()
            img_path = resolve_one(name)
            current_app.logger.info("[AUTH] key=%s name=%s resolved=%s", key, name, img_path)
            if img_path:
                opt_path = optimize_image_for_docx(img_path)
                ctx[placeholder] = InlineImage(tpl, opt_path, width=Mm(155))
            else:
                ctx[placeholder] = ""


    # --- render and save to instance/generated_reports for all other reports ---
    out_dir = os.path.join(current_app.instance_path, "generated_reports")
    os.makedirs(out_dir, exist_ok=True)

    stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"report_{report_id}_{stamp}.docx"
    abs_path = os.path.join(out_dir, filename)

    # ------------------------------------------------------------
    #            PHOTO HANDLING — DOCX DOWNLOAD (FIX)
    # ------------------------------------------------------------
    if paths:
        prepared = prepare_photos(
            tpl,
            paths,
            max_width_mm=155,
            max_height_mm=130
        )

        ctx["photo_pages"] = build_photo_pages(prepared)
        ctx["social_photos"] = [p.image for p in prepared]
    else:
        ctx["photo_pages"] = []
        ctx["social_photos"] = []

    tpl.render(ctx)
    tpl.save(abs_path)

    docx_url = url_for("generated_reports", filename=filename)
    return jsonify({"ok": True, "docx_url": docx_url})


# ===================== PHOTO ID: dedicated endpoints (server only) =====================

import datetime as dt


# --------------------------- PHOTO-ID: Preview (PDF) ---------------------------

def _resolve_photo_id_image(insured_id, report_id, name_or_path):
    """
    Try to resolve the image for the photo-id report.

    - If we get an absolute filesystem path (local Dropbox on the server) – use it.
    - Otherwise, treat it as a file name and look for it under REPORT_MEDIA_DIR
      for this insured (in any subfolder), with a preference to the report folder.
    """
    if not name_or_path:
        return None

    s = str(name_or_path).strip()
    if not s:
        return None

    # If someone accidentally passed a data: URL, we can't handle it here
    if s.startswith("data:"):
        current_app.logger.warning(
            "[PhotoID] got data: URL for insured=%s report=%s; cannot resolve on server",
            insured_id, report_id
        )
        return None

    media_root = current_app.config.get(
        "REPORT_MEDIA_DIR",
        os.path.join(current_app.instance_path, "report_media")
    )

    # 1) Direct filesystem path (Dropbox local path on the server)
    if (s.startswith("/") or ":" in s) and os.path.isfile(s):
        return s

    # 2) If it's a URL, strip querystring/fragment and use the basename
    if "?" in s or "#" in s:
        s = s.split("?", 1)[0].split("#", 1)[0]
    base = os.path.basename(s)

    # Prefer: <media_root>/<insured_id>/<report_id>/<base>
    cand = os.path.join(media_root, str(insured_id), str(report_id), base)
    if os.path.isfile(cand):
        return cand

    # Fallback: search anywhere under this insured's folder
    insured_root = os.path.join(media_root, str(insured_id))
    if os.path.isdir(insured_root):
        for root, _dirs, files in os.walk(insured_root):
            if base in files:
                return os.path.join(root, base)

    current_app.logger.warning(
        "[PhotoID] cannot resolve image %r for insured=%s report=%s",
        name_or_path, insured_id, report_id
    )
    return None


@reports_docx_bp.get("/<int:report_id>/photo-id/preview-pdf")
def photo_id_preview_pdf(report_id: int):
    insured_id = request.args.get("insured_id", type=int) or 0

    # NEW: reference + version from query
    ref_number   = (request.args.get("ref_number") or "").strip()
    reference_no = (request.args.get("reference_no") or "").strip()
    version_no   = request.args.get("version_no", type=int)

    id_date  = request.args.get("id_photo_date_text", "") or request.args.get("id_photo_date", "")
    id_time  = request.args.get("id_photo_time", "")
    id_city  = (request.args.get("id_photo_city") or "").strip()
    id_place = (request.args.get("id_photo_place") or "").strip()


    # Source selected in the "תמונת זיהוי" widget (now usually a short name)
    id_src = (request.args.get("id_photo_src") or "").strip()

    # Also look at selected_photos[] (we always send one basename from the client)
    names = request.args.getlist("selected_photos[]") or request.args.getlist("selected_photos") or []
    base_name = names[0] if names else None

    # If we somehow got the route name 'serve' as a "filename", ignore it
    if base_name and base_name.lower() == "serve":
        base_name = None

    # Prefer the explicit server-side filename; fall back to id_src only if needed
    key = base_name or id_src
    img_path = _resolve_photo_id_image(insured_id, report_id, key)


    current_app.logger.info(
        "[PhotoID][preview] report=%s insured=%s src=%r base=%r -> %r",
        report_id, insured_id, id_src, base_name, img_path
    )

    # build context
    ctx = get_report_context(report_id, insured_id=insured_id)

    ctx = _apply_ref_and_version(
        ctx,
        ref_number=ref_number,
        reference_no=reference_no,
        version_no=version_no,
    )

    # ------------------------------------------------------------
    # Tracking date formatted for Menora header: dd-mm-yyyy
    # ------------------------------------------------------------
    db = ctx.get("db", {})
    ctx["tracking_date_ddmmyyyy"] = _iso_to_ddmmyyyy_dash(
        db.get("tracking_date", "")
    )

    # ------------------------------------------------------------
    # 🔒 RESET PHOTO CONTEXT (prevents leakage between renders)
    # ------------------------------------------------------------
    ctx["photo_pages"] = []
    ctx["social_photos"] = []
    ctx["authorities_table_photo"] = ""
    ctx["authorities_table_photo_2"] = ""


    place_str = ", ".join(v for v in (id_place, id_city) if v)
    ctx.update({
        "id_date":  id_date or _iso_to_dots(ctx.get("ctx", {}).get("activity_date", "")),
        "id_time":  id_time,
        "id_place": place_str,
    })

    template_path = os.path.join(current_app.root_path, "docx_templates", "menora_photo_id.docx")
    current_app.logger.info("[PhotoID][preview] template: %s", template_path)

    tpl = DocxTemplate(template_path)

    if img_path:
        opt_path = optimize_image_for_docx(img_path)
        ctx["id_photo"] = InlineImage(tpl, opt_path, width=Mm(90))
    else:
        ctx["id_photo"] = ""  # keep placeholder empty if not found

    tpl.render(ctx)

    # -> PDF
    buf = io.BytesIO()
    tpl.save(buf)
    buf.seek(0)
    pdf_bytes = _docx_to_pdf_bytes(buf.getvalue())
    return send_file(io.BytesIO(pdf_bytes), mimetype="application/pdf", download_name="preview.pdf")

@reports_docx_bp.post("/<int:report_id>/photo-id/render-docx")
def photo_id_render_docx(report_id: int):
    payload    = request.get_json(silent=True) or {}
    insured_id = payload.get("insured_id") or 0

    # NEW: reference + version from client
    ref_number   = (payload.get("ref_number") or "").strip()
    reference_no = (payload.get("reference_no") or "").strip()
    version_no   = payload.get("version_no")



    id_date  = (payload.get("id_photo_date_text") or payload.get("id_photo_date") or "").strip()
    id_time  = (payload.get("id_photo_time") or "").strip()
    id_city  = (payload.get("id_photo_city") or "").strip()
    id_place = (payload.get("id_photo_place") or "").strip()

    # Source selected in the "תמונת זיהוי" widget (may be full path / URL / data:)
    id_src = (payload.get("id_photo_src") or "").strip()

    # Also consider selected_photos (client sends [basename])
    names = payload.get("selected_photos") or []
    if not names:
        names = request.args.getlist("selected_photos[]") or request.args.getlist("selected_photos") or []

    base_name = names[0] if names else None

    # If we somehow got the route name 'serve' as a "filename", ignore it
    if base_name and base_name.lower() == "serve":
        base_name = None

    # Prefer the explicit server-side filename; fall back to id_src only if needed
    def _normalize_photo_name(v: str) -> str | None:
        if not v:
            return None
        if "name=" in v:
            return v.split("name=", 1)[1].split("&", 1)[0]
        if "?" in v:
            v = v.split("?", 1)[0]
        return os.path.basename(v) or None

    norm_base = _normalize_photo_name(base_name)
    norm_src = _normalize_photo_name(id_src)

    key = norm_base or norm_src
    img_path = _resolve_photo_id_image(insured_id, report_id, key)

    current_app.logger.info(
        "[PhotoID][docx] report=%s insured=%s src=%r base=%r -> %r",
        report_id, insured_id, id_src, base_name, img_path
    )

    ctx = get_report_context(report_id, insured_id=insured_id)

    ctx = _apply_ref_and_version(
        ctx,
        ref_number=ref_number,
        reference_no=reference_no,
        version_no=version_no,
    )

    # ------------------------------------------------------------
    # Tracking date formatted for Menora header: dd-mm-yyyy
    # ------------------------------------------------------------
    db = ctx.get("db", {})
    ctx["tracking_date_ddmmyyyy"] = _iso_to_ddmmyyyy_dash(
        db.get("tracking_date", "")
    )

    # ------------------------------------------------------------
    # 🔒 RESET PHOTO CONTEXT (prevents leakage between renders)
    # ------------------------------------------------------------
    ctx["photo_pages"] = []
    ctx["social_photos"] = []
    ctx["authorities_table_photo"] = ""
    ctx["authorities_table_photo_2"] = ""



    place_str = ", ".join(v for v in (id_place, id_city) if v)
    ctx.update({
        "id_date":  id_date or _iso_to_dots(ctx.get("ctx", {}).get("activity_date", "")),
        "id_time":  id_time,
        "id_place": place_str,
    })

    template_path = os.path.join(current_app.root_path, "docx_templates", "menora_photo_id.docx")
    current_app.logger.info("[PhotoID][docx] template: %s", template_path)

    tpl = DocxTemplate(template_path)

    if img_path:
        opt_path = optimize_image_for_docx(img_path)
        ctx["id_photo"] = InlineImage(tpl, opt_path, width=Mm(90))
    else:
        ctx["id_photo"] = ""  # keep placeholder empty if not found

    tpl.render(ctx)

    # save to instance/generated_reports
    out_dir = os.path.join(current_app.instance_path, "generated_reports")
    os.makedirs(out_dir, exist_ok=True)
    stamp    = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"report_{report_id}_{stamp}.docx"
    abs_path = os.path.join(out_dir, filename)
    tpl.save(abs_path)

    docx_url = url_for("generated_reports", filename=filename)
    current_app.logger.info("[PhotoID][docx] saved: %s -> %s", abs_path, docx_url)

    return jsonify({"ok": True, "docx_url": docx_url})


# ================== /PHOTO ID: dedicated endpoints (server only) =======================

#################### MEDIA UPLOAD  ######################################

@reports_docx_bp.route("/api/insured/<int:insured_id>/media", methods=["GET"])
def api_media_candidates(insured_id: int):
    report_date = (request.args.get("date") or "").strip()
    media_type = (request.args.get("type") or "").strip()  # optional

    from datetime import datetime as _dt

    try:
        dt = _dt.strptime(report_date, "%Y-%m-%d").date()
    except Exception:
        return jsonify({"status": "ok", "media": []})

    q = GilMedia.query.filter(
        GilMedia.insured_id == insured_id,
        GilMedia.taken_date == dt
    )

    if media_type:
        q = q.filter(GilMedia.media_type == media_type)

    rows = q.order_by(GilMedia.taken_at.asc(), GilMedia.media_id.asc()).all()

    out = []
    for r in rows:
        out.append({
            "media_id": r.media_id,
            "media_type": r.media_type,
            "dropbox_path": r.dropbox_path,
            "file_name": r.file_name,
            "taken_at": r.taken_at.isoformat() if r.taken_at else None,
            # UI will use this later; implement thumb route in Step 3 if needed
            # "thumb_url": f"/media/thumb/{r.media_id}"
        })

    return jsonify({"status": "ok", "media": out})


@reports_docx_bp.route("/api/tracking-report/<int:tracking_report_id>/media-links", methods=["GET"])
def api_tracking_report_media_links(tracking_report_id: int):
    links = (GilTrackingReportMedia.query
             .filter(GilTrackingReportMedia.tracking_report_id == tracking_report_id)
             .order_by(GilTrackingReportMedia.sort_order.asc(), GilTrackingReportMedia.id.asc())
             .all())

    media_ids = [l.media_id for l in links]
    media_rows = {}
    if media_ids:
        rows = GilMedia.query.filter(GilMedia.media_id.in_(media_ids)).all()
        media_rows = {r.media_id: r for r in rows}

    out = []
    for l in links:
        r = media_rows.get(l.media_id)
        out.append({
            "media_id": l.media_id,
            "sort_order": l.sort_order,
            "tag": l.tag,
            "media_type": r.media_type if r else None,
            "dropbox_path": r.dropbox_path if r else None,
            "file_name": r.file_name if r else None,
            "taken_at": r.taken_at.isoformat() if (r and r.taken_at) else None,
        })

    return jsonify({"status": "ok", "links": out})


@reports_docx_bp.route("/api/tracking-report/<int:tracking_report_id>/media-links", methods=["POST"])
def api_save_tracking_report_media_links(tracking_report_id: int):
    import json
    user_data = session.get("user")
    user = json.loads(user_data) if user_data else {}
    if not user:
        return jsonify({"status": "error", "message": "Not logged in"}), 401

    payload = request.get_json(silent=True) or {}
    media_ids = payload.get("media_ids") or []

    if not isinstance(media_ids, list):
        return jsonify({"status": "error", "message": "media_ids must be a list"}), 400

    # validate tracking report exists
    rep = GilTrackingReport.query.get(tracking_report_id)
    if not rep:
        return jsonify({"status": "error", "message": "Tracking report not found"}), 404

    # validate media belongs to same insured (important)
    if media_ids:
        cnt = (GilMedia.query
               .filter(GilMedia.media_id.in_(media_ids),
                       GilMedia.insured_id == rep.insured_id)
               .count())
        if cnt != len(set(media_ids)):
            return jsonify({"status": "error", "message": "One or more media items do not belong to this insured"}), 400

    try:
        # delete existing links
        GilTrackingReportMedia.query.filter_by(tracking_report_id=tracking_report_id).delete()

        # insert new links with stable ordering
        user_id = user.get("id")
        for idx, mid in enumerate(media_ids):
            db.session.add(GilTrackingReportMedia(
                tracking_report_id=tracking_report_id,
                media_id=int(mid),
                sort_order=idx,
                created_by_user_id=user_id
            ))

        db.session.commit()
        return jsonify({"status": "ok", "message": "Media links saved", "count": len(media_ids)})

    except Exception as e:
        db.session.rollback()
        current_app.logger.exception("api_save_tracking_report_media_links failed")
        return jsonify({"status": "error", "message": str(e)}), 500















