import io, json, os, uuid, re, subprocess, tempfile
from datetime import datetime, date
from urllib.parse import urljoin, urlparse, parse_qs
from PIL import Image, UnidentifiedImageError  # <— add
from io import BytesIO
from urllib.request import urlopen, Request
from pathlib import Path
from flask import Blueprint, render_template, request, jsonify, current_app, Response
from .report_naming import build_report_display_name
from .report_naming import build_report_filename
from .dropbox_util import get_dbx, build_dropbox_folder_path, ensure_folder_exists, list_pdfs_in_folder



from flask import (
    Blueprint, render_template, request, jsonify, send_file, current_app, url_for
)
from werkzeug.utils import secure_filename

from app import db
from app.models import GilInsured, GilReport

# (only needed if you later embed photos into DOCX, safe to keep)
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


reports_ui_bp = Blueprint('reports_ui', __name__, url_prefix='/reports')


# ---------- helpers ----------

from urllib.parse import urlparse, parse_qs
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _pairs(lst):
    """yield successive pairs: [0:2], [2:4], ..."""
    for i in range(0, len(lst), 2):
        yield lst[i:i+2]


def _classify_photos_by_orientation(urls, resolve_fn):
    """Return (landscape[], portrait[]) using Pillow to decide orientation."""
    lands, ports = [], []
    for url in urls:
        try:
            p = resolve_fn(url)  # absolute local path or None
            if not p:
                # no local path; we’ll guess by file name suffix; fallback = landscape
                (lands if url.lower().endswith(("_p.jpg", "_p.jpeg", "_p.png")) else lands).append(url)
                continue
            with Image.open(p) as im:
                w, h = im.size
            (lands if w >= h else ports).append(url)
        except (UnidentifiedImageError, FileNotFoundError, OSError):
            # If unreadable, default to landscape so it still prints
            lands.append(url)
    return lands, ports


def _build_photos_html(land_urls, port_urls):
    """
    Layout rules:
      • Landscape: max 2 per page, stacked vertically (bigger size).
      • Portrait:  max 2 per page, side-by-side (smaller so both fit).
    """
    css = """
    <style>
      .photo-page { page-break-after: always; }
      .photo-page:last-child { page-break-after: auto; }
      /* keep items intact across pages */
      .photo-block { page-break-inside: avoid; margin: 0 0 8mm 0; }

      /* Landscape — fill width, big height so only 2 fit */
      .land img {
        width: 100%;
        height: auto;
        display: block;
        border: 0;
      }

      /* Portrait — two in a row */
      .row { display: flex; gap: 6mm; }
      .row .port {
        flex: 1 1 0;
      }
      .row .port img {
        width: 100%;
        height: auto;
        display: block;
        border: 0;
      }

      /* Top space to breathe inside section */
      .section-photos { margin-top: 6mm; }
    </style>
    """

    html_parts = [css, '<div class="section-photos">']

    # --- Landscape: two per page, stacked ---
    for pair in _pairs(land_urls):
        html_parts.append('<div class="photo-page">')
        for url in pair:
            html_parts.append(f'''
              <div class="photo-block land">
                <img src="{url}" alt="">
              </div>
            ''')
        html_parts.append('</div>')  # /photo-page

    # --- Portrait: two per page, side-by-side ---
    for pair in _pairs(port_urls):
        html_parts.append('<div class="photo-page">')
        # if single leftover, we still center it in the row
        if len(pair) == 1:
            html_parts.append(f'''
              <div class="photo-block row">
                <div class="port"><img src="{pair[0]}" alt=""></div>
              </div>
            ''')
        else:
            html_parts.append(f'''
              <div class="photo-block row">
                <div class="port"><img src="{pair[0]}" alt=""></div>
                <div class="port"><img src="{pair[1]}" alt=""></div>
              </div>
            ''')
        html_parts.append('</div>')  # /photo-page

    html_parts.append('</div>')  # /section-photos
    return "".join(html_parts)


def _resolve_local_media_path_from_serve_url(url: str) -> str | None:
    """
    Accepts a URL like /reports/photos/serve?case_id=123&report_id=456&name=pic.jpg
    Returns the absolute path under REPORT_MEDIA_DIR, or None if invalid/missing.
    """
    try:
        parsed = urlparse(url)
        q = parse_qs(parsed.query or "")
        case_id   = (q.get("case_id", [""])[0] or "").strip()
        report_id = (q.get("report_id", ["no_report"])[0] or "no_report").strip()
        name      = (q.get("name", [""])[0] or "").strip()
        if not case_id or not name:
            return None
        base_dir = current_app.config.get(
            "REPORT_MEDIA_DIR",
            os.path.join(current_app.instance_path, "report_media")
        )
        p = os.path.abspath(os.path.join(base_dir, case_id, report_id, name))
        # safety: keep inside base_dir
        if not p.startswith(os.path.abspath(base_dir)):
            return None
        return p if os.path.isfile(p) else None
    except Exception:
        return None

def _append_photos_section(document, photo_urls: list[str]):
    """
    Adds a 'תמונות' section with each selected image on its own line.
    Uses width ~5.5 inches to fit A4 portrait margins.
    """
    if not photo_urls:
        return

    document.add_heading('תמונות', level=1)

    for url in photo_urls:
        path = _resolve_local_media_path_from_serve_url(url)
        if not path:
            continue
        try:
            # add picture
            document.add_picture(path, width=Inches(5.5))
            # center the image (applies to the paragraph that holds the picture)
            p = document.paragraphs[-1]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # optional: caption (file name)
            cap = document.add_paragraph(os.path.basename(path))
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            current_app.logger.warning(f"Failed to add picture {path}: {e}")



def _compute_reference(ref_number: str | None, version_no: int) -> str:
    base = (ref_number or "00000").strip()
    return base if (version_no or 0) <= 0 else f"{base}.{version_no}"

def _template_key(insurance: str | None, report_type: str) -> str:
    ins = (insurance or '').strip() or 'כללי'
    return f"{ins}_{report_type}".lower()

def _he_date(iso_date: str | None) -> str:
    """YYYY-MM-DD -> '30 בספטמבר 2025' (or today if missing/invalid)."""
    if not iso_date:
        dt = date.today()
    else:
        try:
            y, m, d = map(int, iso_date.split('-'))
            dt = date(y, m, d)
        except Exception:
            dt = date.today()
    months = ["ינואר","פברואר","מרץ","אפריל","מאי","יוני","יולי","אוגוסט","ספטמבר","אוקטובר","נובמבר","דצמבר"]
    return f"{dt.day} ב{months[dt.month-1]} {dt.year}"

def _ddmmyyyy(iso_date: str | None) -> str:
    """YYYY-MM-DD -> DD.MM.YYYY (empty if missing/invalid)."""
    try:
        y, m, d = map(int, (iso_date or '').split('-'))
        return f"{d:02d}.{m:02d}.{y}"
    except Exception:
        return ""

def _get_first_nonempty(obj, *names, default=""):
    for n in names:
        if hasattr(obj, n):
            v = getattr(obj, n)
            if v not in (None, "", "NULL"):
                return v
    return default


def _wkhtmltopdf_bytes(body_html: str, header_html: str, footer_html: str) -> bytes:
    """Render HTML -> PDF bytes with wkhtmltopdf (header/footer HTML enabled)."""
    wkhtml = current_app.config.get("WKHTMLTOPDF_CMD", "/usr/bin/wkhtmltopdf")
    if not os.path.exists(wkhtml):
        raise RuntimeError(f"wkhtmltopdf not found at: {wkhtml!r}")

    # Write three temp HTML files
    with tempfile.NamedTemporaryFile(suffix=".html", delete=False) as fbody, \
         tempfile.NamedTemporaryFile(suffix=".html", delete=False) as fhead, \
         tempfile.NamedTemporaryFile(suffix=".html", delete=False) as ffoot:
        fbody.write(body_html.encode("utf-8"));  fbody.flush()
        fhead.write(header_html.encode("utf-8")); fhead.flush()
        ffoot.write(footer_html.encode("utf-8")); ffoot.flush()
        body_path, head_path, foot_path = fbody.name, fhead.name, ffoot.name

    try:
        # Important:
        # - Use units (e.g., 'mm') for margins
        # - Give enough top/bottom margin so header/footer can render
        # - Spacing controls the gap between the page edge margin and the header/footer content
        cmd = [
            wkhtml,
            "--quiet",
            "--enable-local-file-access",
            "--dpi", "150",

            "--margin-top",    "20mm",
            "--header-html",    head_path,
            "--header-spacing", "3",

            "--margin-bottom", "16mm",
            "--footer-html",    foot_path,
            "--footer-spacing", "2",

            body_path, "-"  # input and stdout
        ]

        proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        if proc.returncode != 0:
            raise RuntimeError(proc.stderr.decode("utf-8", "ignore"))
        return proc.stdout
    finally:
        for p in (body_path, head_path, foot_path):
            try:
                os.remove(p)
            except Exception:
                pass
# --- Detect insurer & claim type profile ---
def get_report_profile(insurer_name, claim_type):
    # Menora Life (אובדן כושר עבודה)
    if insurer_name == "מנורה" and claim_type == "אכע":
        return {
            "default_type": "menora_life_followup",
            "report_types": [
                {"value": "menora_life_followup", "label": "דוח מעקב ביטוח חיים"},
                {"value": "menora_life_photos",   "label": "דוח תמונות"},
                {"value": "menora_life_photoid",  "label": "תמונת זיהוי"},
                {"value": "menora_life_invoice",  "label": "חשבונית עסקה"},
            ]
        }
    return None

# ---------- UI pages ----------

@reports_ui_bp.route("/editor", methods=["GET"])
def editor_page():
    """
    Render the report editor (main UI) with proper report type options
    depending on insurer and claim type.
    """

    # Pre-selected ids (when coming from admin_insured etc.)
    insured_id = request.args.get("insured_id", type=int)
    report_id  = request.args.get("report_id", type=int)

    # Full list for the dropdown on the left
    insureds = (
        db.session.query(GilInsured)
        .order_by(GilInsured.last_name.asc())
        .all()
    )

    # The currently-selected insured (if any)
    insured = None
    if insured_id:
        insured = next((i for i in insureds if i.id == insured_id), None)

    # --- Build report type profile (Menora / אכע → life reports) ---
    if insured:
        # NOTE: model field is "insurance", not "insurer_name"
        profile = get_report_profile(insured.insurance, insured.claim_type)
    else:
        profile = None

    if profile:
        report_types = profile["report_types"]          # already [{value,label}, ...]
        selected_report_type = profile["default_type"]  # e.g. "menora_life_followup"
    else:
        # Default bundle for סיעודי / everything that is NOT Menora אכע
        report_types = [
            {"value": "SIUDI", "label": "דוח חקירה סיעודי"},
            {"value": "ID_PHOTOS", "label": "תמונת זיהוי"},
            {"value": "SIUDI_INVOICE", "label": "חשבונית סיעודי"},
        ]
        selected_report_type = "SIUDI"

    ctx = {
        "insureds":            insureds,            # <-- needed for the dropdown
        "insured":             insured,
        "insured_id":          insured_id,
        "report_id":           report_id,
        "report_types":        report_types,
        "selected_report_type": selected_report_type,
    }

    current_app.logger.info(
        "[EDITOR] insurer=%s claim_type=%s -> default=%s",
        getattr(insured, "insurance", None),
        getattr(insured, "claim_type", None),
        selected_report_type,
    )

    return render_template("reports/editor.html", **ctx)


@reports_ui_bp.route('/save_draft', methods=['POST'])
def save_draft():
    payload = request.get_json(silent=True) or {}

    report_id = payload.get('report_id')
    insured_id = payload.get('insured_id')
    report_type = (payload.get('report_type') or 'TRACKING').strip()

    insured = db.session.get(GilInsured, insured_id) if insured_id else None

    if report_id:
        rpt = db.session.get(GilReport, int(report_id))
        if not rpt:
            return jsonify({'status': 'error', 'message': 'Report not found'}), 404

        if not insured and rpt.case_id:
            insured = db.session.get(GilInsured, rpt.case_id)

    else:
        if not insured:
            return jsonify({'status': 'error', 'message': 'Missing insured_id'}), 400

        rpt = GilReport(
            case_id=insured.id,
            report_type=report_type,
            template_key=_template_key(getattr(insured, 'insurance', None), report_type),
            status='Draft',
            version_no=int(payload.get('version_no') or 0),
        )
        db.session.add(rpt)

    if getattr(rpt, 'status', None) != 'Final':
        rpt.status = 'Draft'

    editor_json = payload
    rpt.editor_json = json.dumps(editor_json, ensure_ascii=False)

    ref_number = _get_first_nonempty(insured or object(), 'ref_number', 'ref', default=None)
    rpt.reference_no = _compute_reference(ref_number, int(rpt.version_no or 0))

    # ---- New naming convention ----
    full_name = (
        (payload.get('db') or {}).get('full_name')
        or getattr(insured, 'full_name', None)
        or getattr(insured, 'name', None)
        or ''
    )

    invoice_no = (
        payload.get('inv_number')
        or (payload.get('invoice') or {}).get('inv_number')
        or ''
    )

    rpt.title = build_report_display_name(
        report_type=report_type,
        full_name=full_name,
        reference_no=rpt.reference_no or '',
        invoice_no=invoice_no,
    )

    rpt.updated_at = datetime.utcnow()
    db.session.commit()

    return jsonify({
        'status': 'ok',
        'report_id': rpt.id,
        'status_str': rpt.status,
        'version_no': rpt.version_no,
        'reference_no': rpt.reference_no,
        'title': rpt.title,
    })

@reports_ui_bp.route('/load_draft', methods=['GET'])
def load_draft():
    insured_id = request.args.get('insured_id', type=int)
    report_id = request.args.get('report_id', type=int)

    rpt = None

    if report_id:
        rpt = db.session.get(GilReport, report_id)
        if not rpt:
            return jsonify({'status': 'error', 'message': 'Report not found'}), 404

    elif insured_id:
        rpt = (
            GilReport.query
            .filter(
                GilReport.case_id == insured_id,
                GilReport.status == 'Draft'
            )
            .order_by(GilReport.updated_at.desc(), GilReport.id.desc())
            .first()
        )
        if not rpt:
            return jsonify({'status': 'empty', 'message': 'No draft found'})

    else:
        return jsonify({'status': 'error', 'message': 'Missing insured_id or report_id'}), 400

    try:
        editor_data = json.loads(rpt.editor_json) if rpt.editor_json else {}
    except Exception:
        editor_data = {}

    return jsonify({
        'status': 'ok',
        'report_id': rpt.id,
        'status_str': rpt.status,
        'version_no': rpt.version_no,
        'reference_no': rpt.reference_no,
        'report_type': rpt.report_type,
        'editor_json': editor_data
    })

@reports_ui_bp.route('/finalize', methods=['POST'])
def finalize():
    from dropbox.files import WriteMode

    payload = request.form or request.args or {}
    pdf_file = request.files.get("pdf_file")

    current_app.logger.info(f"[REPORTS_UI FINALIZE] content_type={request.content_type}")
    current_app.logger.info(f"[REPORTS_UI FINALIZE] form_keys={list(payload.keys()) if payload else []}")
    current_app.logger.info(f"[REPORTS_UI FINALIZE] has_pdf_file={bool(pdf_file)}")

    action = payload.get('action') or 'finalize'
    report_id = payload.get('report_id')

    if not report_id:
        return jsonify({'status': 'error', 'message': 'Missing report_id'}), 400

    rpt = db.session.get(GilReport, int(report_id))
    if not rpt:
        return jsonify({'status': 'error', 'message': 'Report not found'}), 404

    insured = db.session.get(GilInsured, rpt.case_id)
    if not insured:
        return jsonify({'status': 'error', 'message': 'Insured not found'}), 404

    ref_number = _get_first_nonempty(insured, 'ref_number', 'ref', default=None)

    try:
        if action == 'version':
            rpt.version_no = (rpt.version_no or 0) + 1
            rpt.status = 'Revised'
            rpt.reference_no = _compute_reference(ref_number, rpt.version_no)

            db.session.commit()
            return jsonify({
                'status': 'ok',
                'status_str': rpt.status,
                'version_no': rpt.version_no,
                'reference_no': rpt.reference_no
            })

        elif action == 'send_to_insurer':
            rpt.status = 'Submitted'

            db.session.commit()
            return jsonify({
                'status': 'ok',
                'status_str': rpt.status,
                'version_no': rpt.version_no,
                'reference_no': rpt.reference_no
            })

        elif action == 'finalize':
            if not pdf_file:
                return jsonify({'status': 'error', 'message': 'Missing PDF file'}), 400

            rpt.reference_no = _compute_reference(ref_number, int(rpt.version_no or 0))

            full_name = (
                _get_first_nonempty(insured, 'full_name')
                or f"{_get_first_nonempty(insured, 'last_name')} {_get_first_nonempty(insured, 'first_name')}".strip()
                or _get_first_nonempty(insured, 'name')
                or ""
            )

            final_pdf_filename = build_report_filename(
                report_type=(payload.get("template") or rpt.template_key or rpt.report_type or "tracking"),
                full_name=full_name,
                reference_no=(payload.get("reference_no") or rpt.reference_no or ""),
                invoice_no="",
                ext="pdf",
                version_no=rpt.version_no
            )

            case_root = build_dropbox_folder_path(
                insurance=_get_first_nonempty(insured, 'insurance'),
                claim_type=_get_first_nonempty(insured, 'claim_type'),
                last_name=_get_first_nonempty(insured, 'last_name'),
                first_name=_get_first_nonempty(insured, 'first_name'),
                id_number=_get_first_nonempty(insured, 'id_number'),
                claim_number=_get_first_nonempty(insured, 'claim_number'),
            )

            if not case_root:
                return jsonify({'status': 'error', 'message': 'Could not build Dropbox insured folder path'}), 400

            dropbox_reports_dir = f"{case_root}/דוחות"
            dbx = get_dbx()
            ensure_folder_exists(dbx, dropbox_reports_dir)

            target_dropbox_path = f"{dropbox_reports_dir}/{final_pdf_filename}"

            pdf_bytes = pdf_file.read()
            if not pdf_bytes:
                return jsonify({'status': 'error', 'message': 'Uploaded PDF is empty'}), 400

            dbx.files_upload(
                pdf_bytes,
                target_dropbox_path,
                mode=WriteMode.overwrite,
                mute=True
            )

            rpt.status = 'Final'

            if hasattr(rpt, 'dropbox_path'):
                rpt.dropbox_path = target_dropbox_path
            if hasattr(rpt, 'generated_pdf_path'):
                rpt.generated_pdf_path = target_dropbox_path
            if hasattr(rpt, 'final_pdf_path'):
                rpt.final_pdf_path = target_dropbox_path
            if hasattr(rpt, 'title'):
                rpt.title = final_pdf_filename.rsplit('.', 1)[0]

            db.session.commit()

            return jsonify({
                'status': 'ok',
                'status_str': rpt.status,
                'version_no': rpt.version_no,
                'reference_no': rpt.reference_no,
                'dropbox_path': target_dropbox_path
            })

        return jsonify({'status': 'error', 'message': f'Unsupported action: {action}'}), 400

    except Exception as e:
        db.session.rollback()
        current_app.logger.exception("reports_ui.finalize failed")
        return jsonify({'status': 'error', 'message': str(e)}), 500


# ---------- Preview (header/footer via wkhtmltopdf) ----------

@reports_ui_bp.route('/preview', methods=['POST'])
def preview():
    payload     = request.get_json(silent=True) or {}
    insured_id  = payload.get('insured_id')
    report_type = payload.get('report_type', 'TRACKING')
    photos      = payload.get('photos') or []

    insured = db.session.get(GilInsured, insured_id) if insured_id else None
    if not insured:
        return jsonify({'error': 'missing insured'}), 400

    insurer     = (insured.insurance or "").strip()
    claim_type  = (insured.claim_type or "").strip()

    # refs & dates
    reference_no     = payload.get('reference_no') or _get_first_nonempty(insured, 'ref_number','ref', default="00000")
    report_date_text = _he_date(payload.get('report_date'))
    activity_date_fmt= _ddmmyyyy(payload.get('activity_date'))

    # body fields
    full_name    = f"{_get_first_nonempty(insured, 'last_name')} {_get_first_nonempty(insured, 'first_name')}".strip()
    address      = f"{_get_first_nonempty(insured, 'city')} {_get_first_nonempty(insured, 'address')}".strip()
    id_number    = _get_first_nonempty(insured, 'id_number')
    claim_number = _get_first_nonempty(insured, 'claim_number', 'claim_no', 'claim', 'claimnumber')
    injury_type  = "מצב סיעודי"

    birth_year = ""
    try:
        bd = insured.birth_date
        if bd:
            birth_year = bd[:4] if isinstance(bd, str) else str(bd.year)
    except Exception:
        pass

    # ---------- Static URLs (HTTP, not file://) ----------
    base_url = request.url_root.rstrip('/')
    header_url = url_for('static', filename='report_header_gil.png', _external=True)
    footer_url = url_for('static', filename='report_footer_gil.png', _external=True)

    # ---------- Debug: verify static folder & files ----------
    current_app.logger.info("STATIC folder: %s", current_app.static_folder)
    header_path = os.path.join(current_app.static_folder, 'report_header_gil.png')
    footer_path = os.path.join(current_app.static_folder, 'report_footer_gil.png')
    current_app.logger.info("Header file exists: %s", os.path.exists(header_path))
    current_app.logger.info("Footer file exists: %s", os.path.exists(footer_path))
    current_app.logger.info("Header URL: %s", header_url)
    current_app.logger.info("Footer URL: %s", footer_url)

    # Optional: probe only for http(s) URLs (requests can’t fetch file://)
    def _is_http(u: str) -> bool:
        return u.startswith('http://') or u.startswith('https://')

    try:
        if _is_http(header_url):
            import requests
            r = requests.get(header_url, timeout=3)
            current_app.logger.info("Header URL HTTP %s (bytes=%s)", r.status_code, len(r.content))
        if _is_http(footer_url):
            import requests
            r = requests.get(footer_url, timeout=3)
            current_app.logger.info("Footer URL HTTP %s (bytes=%s)", r.status_code, len(r.content))
    except Exception as e:
        current_app.logger.warning("Static URL probe failed: %s", e)

    # ----------------- Body HTML -----------------
    if insurer == "מנורה" and "סיעוד" in claim_type and report_type == "TRACKING":
        body_html = render_template(
            "reports/templates/menora_siudi_tracking.html",
            full_name=full_name, address=address,
            id_number=id_number, claim_number=claim_number,
            injury_type=injury_type, birth_year=birth_year,
            activity_date=activity_date_fmt,
            photos=photos,
        )
    else:
        photos_css = """
        <style>
          .photos h3 { margin-top: 28px; }
          .photo { page-break-inside: avoid; margin: 10px 0 20px; text-align:center; }
          .photo img { max-width: 100%; width: 520px; height: auto; display:inline-block; }
          .caption { font-size: 12px; color:#666; margin-top: 6px; }
        </style>
        """
        photos_html = ""
        if photos:
            items = []
            for u in photos:
                name = u.split("name=")[-1] if "name=" in u else ""
                items.append(f"""
                <div class="photo">
                  <img src="{u}" alt="">
                  <div class="caption">{name}</div>
                </div>
                """)
            photos_html = f'<div class="photos"><h3>תמונות</h3>{"".join(items)}</div>'

        body_html = f"""<!doctype html><meta charset="utf-8">
        <body dir="rtl" style="font-family:Arial;padding:40px">
          {photos_css}
          <h3>תצוגה מקדימה</h3>
          <p>טרם קיימת תבנית לדגם זה.</p>
          <p><b>חברה:</b> {insurer} · <b>סוג תביעה:</b> {claim_type} · <b>סוג דו"ח:</b> {report_type}</p>
          {photos_html}
        </body>"""

    # ----------------- Header & Footer HTML -----------------
    header_html = f"""<!doctype html><html lang="he" dir="rtl"><meta charset="utf-8">
    <style>
      html,body{{margin:0;padding:0}}
      .strip img{{width:100%; height:72px; display:block; object-fit:cover}}
    </style>
    <body>
      <div class="strip"><img src="{header_url}" alt=""></div>
    </body></html>"""

    footer_html = f"""<!doctype html><html lang="he" dir="rtl"><meta charset="utf-8">
    <style>
      html,body{{margin:0;padding:0;font-family:'Assistant',Arial,sans-serif;}}
      .foot{{position:relative; height:50px; box-sizing:border-box;}}
      .pageno{{
        position:absolute; left:12mm; bottom:2px;
        width:12mm; height:12mm; background:#333; color:#fff; border-radius:3px;
      }}
      .pageno .page{{
        position:absolute; left:50%; top:50%; transform:translate(-50%,-50%);
        font-size:10pt; font-weight:700; line-height:1; margin:0; padding:0;
      }}
      .brand{{ position:absolute; right:12mm; bottom:0; }}
      .brand img{{ height:40px; width:auto; display:block; }}
    </style>
    <body onload="subst()">
      <div class="foot">
        <div class="pageno"><span class="page"></span></div>
        <div class="brand"><img src="{footer_url}" alt=""></div>
      </div>
      <script>
      function subst(){{
        var vars={{}}, q=window.location.search.substring(1).split('&');
        for (var i=0;i<q.length;i++) {{ var p=q[i].split('=',2); vars[p[0]]=decodeURIComponent(p[1]||''); }}
        var els=document.getElementsByClassName('page');
        for (var j=0;j<els.length;j++) els[j].textContent = vars.page || '';
      }}
      </script>
    </body></html>"""

    # ----------------- Absolutize URLs -----------------
    def absolutize_urls(html: str) -> str:
        if not html:
            return html
        if '<base ' not in html:
            if '<head>' in html:
                html = html.replace('<head>', f'<head><base href="{base_url}/">', 1)
            else:
                html = f'<!doctype html><head><base href="{base_url}/"></head>' + html
        html = re.sub(r'((?:src|href)=["\'])(/[^"\']*)', rf'\1{base_url}\2', html)
        return html

    body_html   = absolutize_urls(body_html)
    header_html = absolutize_urls(header_html)
    footer_html = absolutize_urls(footer_html)

    # ----------------- Center the table after הנדון -----------------
    extra_css = """
    <style>
      h1 + table, h1 + div > table { margin: 0 auto; }
      h1 + table td, h1 + div > table td { padding: 0 12px; font-size: 16pt; }
    </style>
    """
    m = re.search(r"(?i)<body[^>]*>", body_html)
    body_html = body_html[:m.end()] + extra_css + body_html[m.end():] if m else extra_css + body_html

    # ----------------- Letter top (single injection) -----------------
    recipient_insurer = insurer or "מנורה – חברה לביטוח"
    recipient_dept    = "מחלקת תביעות סיעודי"
    letter_css = """
    <style>
      .letter-top{ display:flex; justify-content:space-between; align-items:flex-start;
                   margin: 6px 0 12px 0; font-family: 'Assistant', Arial, sans-serif; }
      .lt-right{ text-align:right; line-height:1.2; }
      .lt-right .line-2, .lt-right .line-3{ font-weight:800; }
      .lt-left{ text-align:left; font-weight:800; }
    </style>
    """
    letter_top_html = f"""
    {letter_css}
    <div class="letter-top" dir="rtl">
      <div class="lt-right">
        <div>לכבוד</div>
        <div class="line-2">{recipient_insurer}</div>
        <div class="line-3">{recipient_dept}</div>
      </div>
      <div class="lt-left">
        <div>{report_date_text}</div>
        <div>מספרנו: {reference_no}</div>
      </div>
    </div>
    """
    if 'class="letter-top"' not in body_html:
        m = re.search(r"(?i)<body[^>]*>", body_html)
        body_html = body_html[:m.end()] + letter_top_html + body_html[m.end():] if m else letter_top_html + body_html

    # ----------------- Photos layout (unchanged) -----------------
    def _pairs(seq):
        for i in range(0, len(seq), 2):
            yield seq[i:i+2]

    def _classify(links):
        try:
            from PIL import Image, UnidentifiedImageError
        except Exception:
            Image = None
            UnidentifiedImageError = Exception

        lands, ports = [], []
        for u in links:
            if not Image:
                lands.append(u); continue
            try:
                req = Request(u, headers={"User-Agent": "PreviewClassifier/1.0"})
                with urlopen(req, timeout=5) as resp:
                    data = resp.read()
                with Image.open(BytesIO(data)) as im:
                    w, h = im.size
                (lands if w >= h else ports).append(u)
            except (UnidentifiedImageError, OSError, ValueError, TimeoutError, Exception):
                lands.append(u)
        return lands, ports

    def _photos_html(land_urls, port_urls):
        css = """
        <style>
          .section-photos { margin-top: 6mm; }
          .photo-page { page-break-after: always; }
          .photo-page:last-child { page-break-after: auto; }
          .photo-block { page-break-inside: avoid; margin: 0 0 8mm 0; }
          .land img { width: 100%; height: auto; display:block; border:0; }
          .row { display:flex; gap: 6mm; }
          .row .port { flex: 1 1 0; }
          .row .port img { width:100%; height:auto; display:block; border:0; }
        </style>
        """
        parts = [css, '<div class="section-photos">']
        for pair in _pairs(land_urls):
            parts.append('<div class="photo-page">')
            for u in pair:
                parts.append(f'<div class="photo-block land"><img src="{u}" alt=""></div>')
            parts.append('</div>')
        for pair in _pairs(port_urls):
            parts.append('<div class="photo-page">')
            if len(pair) == 1:
                parts.append(f'<div class="photo-block row"><div class="port"><img src="{pair[0]}" alt=""></div></div>')
            else:
                parts.append(f'''
                  <div class="photo-block row">
                    <div class="port"><img src="{pair[0]}" alt=""></div>
                    <div class="port"><img src="{pair[1]}" alt=""></div>
                  </div>
                ''')
            parts.append('</div>')
        parts.append('</div>')
        return "".join(parts)

    if photos:
        land, port = _classify(photos)
        gallery = _photos_html(land, port)
        placeholder = "(יישום חלק התמונות יתווסף בהמשך)"
        if placeholder in body_html:
            body_html = body_html.replace(placeholder, gallery)
        else:
            body_html = re.sub(r"\(?יישום חלק התמונות יתווסף בהמשך\)?", gallery, body_html)

    # ----------------- PDF -----------------
    pdf_bytes = _wkhtmltopdf_bytes(body_html, header_html, footer_html)
    return send_file(io.BytesIO(pdf_bytes), mimetype='application/pdf', as_attachment=False, download_name='preview.pdf')




# --- Add these imports at top of reports_ui.py ---
import os, uuid
from flask import current_app
from werkzeug.utils import secure_filename
from app.models import GilInsured
from app.dropbox_util import dbx, build_dropbox_folder_path, list_case_images



ALLOWED_EXTS = {"jpg", "jpeg", "png"}

def _uploads_root():
    root = os.path.join(current_app.root_path, "static", "uploads", "reports")
    os.makedirs(root, exist_ok=True)
    return root

def _public_url_for(rel_path: str) -> str:
    # rel_path like 'uploads/reports/abcd/pic.jpg'
    rel_norm = rel_path.replace("\\", "/")  # <- do the replace outside the f-string (Py 3.11 safe)
    return f"/static/{rel_norm}"


# app/reports_ui.py (only the dropbox branch inside list_photos)
@reports_ui_bp.route('/photos/list', methods=['POST'])
def list_photos():
    data = request.get_json(silent=True) or {}
    source = (data.get('source') or '').strip()
    images = []

    if source == "dropbox":
        insured_id = data.get("insured_id")
        if not insured_id:
            return jsonify({"images": []})

        insured = db.session.get(GilInsured, int(insured_id))
        if not insured:
            return jsonify({"images": []})

        folder = build_dropbox_folder_path(
            insured.insurance or "",
            insured.claim_type or "",
            insured.last_name or "",
            insured.first_name or "",
            insured.id_number or "",
            insured.claim_number or "",
        )
        if not folder:
            return jsonify({"images": []})

        images = list_case_images(dbx, folder)  # returns thumbnails (fast)
        if not images:
            return jsonify({"images": []})

        return jsonify({"images": images})

    # … handle 'local' branch if you keep it …
    return jsonify({"images": []})


@reports_ui_bp.route('/photos/upload', methods=['POST'])
def upload_photos():
    """
    FormData: files[] (multiple), local_token
    Saves under /static/uploads/reports/<token>/filename and returns URLs.
    """
    token = request.form.get("local_token") or str(uuid.uuid4())
    files = request.files.getlist("files[]") or request.files.getlist("files")
    saved = []

    if not files:
        return jsonify({"token": token, "images": []})

    session_dir = os.path.join(_uploads_root(), token)
    os.makedirs(session_dir, exist_ok=True)

    for f in files:
        if not f or not f.filename:
            continue
        ext = f.filename.rsplit(".", 1)[-1].lower()
        if ext not in ALLOWED_EXTS:
            continue
        safe = secure_filename(f.filename)
        # avoid collisions
        base, extn = os.path.splitext(safe)
        final = f"{base}_{uuid.uuid4().hex[:6]}{extn}"
        out_path = os.path.join(session_dir, final)
        f.save(out_path)
        rel = os.path.join("uploads", "reports", token, final)
        saved.append({"name": final, "url": _public_url_for(rel)})

    return jsonify({"token": token, "images": saved})


# Adjust to your structure
ALLOWED_EXTENSIONS = {"jpg","jpeg","png","gif","webp","bmp","tif","tiff","mp4","mov","avi","mkv"}

def allowed_file(fn: str) -> bool:
    return "." in fn and fn.rsplit(".",1)[1].lower() in ALLOWED_EXTENSIONS

def ensure_dir(path):
    os.makedirs(path, exist_ok=True)
    return path


@reports_ui_bp.post("/import_local_dropbox")
def import_local_dropbox():
    """
    Accept both:
      - single file under form key 'file'
      - multiple files under 'files[]' or 'files'
    Saves into REPORT_MEDIA_DIR/<case_id>/<report_id>/ and returns list of saved items.
    """
    from werkzeug.utils import secure_filename

    # collect files (multi or single)
    files = (
        request.files.getlist("files[]")
        or request.files.getlist("files")
        or ([request.files.get("file")] if request.files.get("file") else [])
    )
    if not files:
        return jsonify({"status": "error", "message": "no files"}), 400

    allowed = {"jpg","jpeg","png","gif","webp","bmp","tif","tiff","mp4","mov","avi","mkv","webm"}
    base_dir = current_app.config.get(
        "REPORT_MEDIA_DIR",
        os.path.join(current_app.instance_path, "report_media")
    )
    case_id   = (request.form.get("case_id") or "no_case").strip()
    report_id = (request.form.get("report_id") or "no_report").strip()

    target_dir = os.path.join(base_dir, case_id, report_id)
    os.makedirs(target_dir, exist_ok=True)

    saved = []
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    for f in files:
        if not f or not f.filename:
            continue
        ext = f.filename.rsplit(".", 1)[-1].lower()
        if ext not in allowed:
            continue

        safe = secure_filename(f.filename)
        name, ext2 = os.path.splitext(safe)
        final = f"{name}_{ts}_{uuid.uuid4().hex[:6]}{ext2.lower()}"
        path = os.path.join(target_dir, final)
        f.save(path)

        # URL to serve later via /reports/photos/serve
        serve_url = (
            f"/reports/photos/serve?case_id={case_id}"
            f"&report_id={report_id}&name={final}"
        )
        saved.append({"name": final, "url": serve_url})

    return jsonify({"status": "ok", "saved": saved})


# ------- list uploaded photos for this report (from REPORT_MEDIA_DIR) -------
@reports_ui_bp.post("/photos/list_report")
def list_report_photos():
    from mimetypes import guess_type
    data = request.get_json(silent=True) or {}
    case_id   = (data.get("case_id") or "").strip()
    report_id = (data.get("report_id") or "").strip()
    if not case_id:
        return jsonify({"images": []})

    base_dir = current_app.config.get(
        "REPORT_MEDIA_DIR",
        os.path.join(current_app.instance_path, "report_media")
    )
    # If report_id is empty, look under a default folder so we still show things
    report_id = report_id or "no_report"
    folder = os.path.join(base_dir, case_id, report_id)
    if not os.path.isdir(folder):
        return jsonify({"images": []})

    allowed = {"jpg","jpeg","png","gif","webp","bmp","tif","tiff"}
    images = []
    for name in sorted(os.listdir(folder)):
        ext = name.rsplit(".", 1)[-1].lower()
        if ext not in allowed:
            continue
        url = (
            f"/reports/photos/serve"
            f"?case_id={case_id}&report_id={report_id}&name={name}"
        )
        images.append({"name": name, "url": url})  # your renderThumb will load it
    return jsonify({"images": images})


# ------- serve a single uploaded photo (send_file) -------
@reports_ui_bp.get("/photos/serve")
def serve_report_photo():
    case_id   = (request.args.get("case_id") or "").strip()
    report_id = (request.args.get("report_id") or "no_report").strip()
    name      = (request.args.get("name") or "").strip()
    if not case_id or not name:
        return ("", 404)

    base_dir = current_app.config.get(
        "REPORT_MEDIA_DIR",
        os.path.join(current_app.instance_path, "report_media")
    )
    path = os.path.join(base_dir, case_id, report_id, name)

    # basic path safety
    base_dir_abs = os.path.abspath(base_dir)
    path_abs = os.path.abspath(path)
    if not path_abs.startswith(base_dir_abs) or not os.path.isfile(path_abs):
        return ("", 404)

    return send_file(path_abs)

@reports_ui_bp.route('/manual-finalize/list-pdfs', methods=['GET'])
def manual_finalize_list_pdfs():
    report_id = request.args.get("report_id")
    if not report_id:
        return jsonify({"status": "error", "message": "Missing report_id"}), 400

    rpt = db.session.get(GilReport, int(report_id))
    if not rpt:
        return jsonify({"status": "error", "message": "Report not found"}), 404

    insured = db.session.get(GilInsured, rpt.case_id)
    if not insured:
        return jsonify({"status": "error", "message": "Insured not found"}), 404

    case_root = build_dropbox_folder_path(
        insurance=_get_first_nonempty(insured, 'insurance'),
        claim_type=_get_first_nonempty(insured, 'claim_type'),
        last_name=_get_first_nonempty(insured, 'last_name'),
        first_name=_get_first_nonempty(insured, 'first_name'),
        id_number=_get_first_nonempty(insured, 'id_number'),
        claim_number=_get_first_nonempty(insured, 'claim_number'),
    )

    if not case_root:
        return jsonify({"status": "error", "message": "Could not build Dropbox insured folder path"}), 400

    reports_dir = f"{case_root}/דוחות"
    dbx = get_dbx()
    pdfs = list_pdfs_in_folder(dbx, reports_dir)

    return jsonify({
        "status": "ok",
        "report_id": rpt.id,
        "folder": reports_dir,
        "files": pdfs
    })

@reports_ui_bp.route('/manual-finalize/select', methods=['POST'])
def manual_finalize_select():
    payload = request.get_json(silent=True) or request.form or {}

    report_id = payload.get("report_id")
    selected_path = (payload.get("dropbox_path") or "").strip()
    selected_name = (payload.get("file_name") or "").strip()

    if not report_id:
        return jsonify({"status": "error", "message": "Missing report_id"}), 400
    if not selected_path:
        return jsonify({"status": "error", "message": "Missing dropbox_path"}), 400

    rpt = db.session.get(GilReport, int(report_id))
    if not rpt:
        return jsonify({"status": "error", "message": "Report not found"}), 404

    try:
        rpt.status = "Final"
        rpt.finalize_mode = "manual"
        rpt.generated_pdf_path = selected_path

        if hasattr(rpt, "title") and selected_name:
            rpt.title = selected_name.rsplit(".", 1)[0]

        db.session.commit()

        return jsonify({
            "status": "ok",
            "status_str": rpt.status,
            "finalize_mode": rpt.finalize_mode,
            "generated_pdf_path": rpt.generated_pdf_path,
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.exception("manual_finalize_select failed")
        return jsonify({"status": "error", "message": str(e)}), 500

@reports_ui_bp.route("/dropbox/open-pdf")
def open_dropbox_pdf():
    path = request.args.get("path")

    if not path:
        return "Missing path", 400

    try:
        dbx = get_dbx()

        md, res = dbx.files_download(path)
        pdf_bytes = res.content

        return Response(
            pdf_bytes,
            mimetype="application/pdf"
        )

    except Exception as e:
        current_app.logger.exception("open_dropbox_pdf failed")
        return str(e), 500

@reports_ui_bp.route('/create-version', methods=['POST'])
def create_version():
    payload = request.get_json(silent=True) or request.form or {}

    source_report_id = payload.get("source_report_id")
    version_no = payload.get("version_no")

    if not source_report_id:
        return jsonify({"status": "error", "message": "Missing source_report_id"}), 400

    try:
        version_no = int(version_no)
    except Exception:
        return jsonify({"status": "error", "message": "Invalid version_no"}), 400

    if version_no < 0:
        return jsonify({"status": "error", "message": "version_no must be 0 or higher"}), 400

    src = db.session.get(GilReport, int(source_report_id))
    if not src:
        return jsonify({"status": "error", "message": "Source report not found"}), 404

    # Prevent duplicate version per case + reference + template
    existing = (
        GilReport.query
        .filter(
            GilReport.case_id == src.case_id,
            GilReport.template_key == src.template_key,
            GilReport.reference_no == src.reference_no,
            GilReport.version_no == version_no
        )
        .first()
    )
    if existing:
        return jsonify({
            "status": "error",
            "message": f"Version {version_no} already exists for this report"
        }), 400

    try:
        new_rpt = GilReport(
            case_id=src.case_id,
            insurer_id=src.insurer_id,
            report_type=src.report_type,
            template_key=src.template_key,
            title=src.title,
            status='Draft',
            version_no=version_no,
            reference_no=src.reference_no,
            editor_json=src.editor_json,
            generated_html=None,
            generated_pdf_path=None,
            created_by=src.created_by,
            updated_by=src.updated_by,
        )

        # optional column if mapped in model
        if hasattr(new_rpt, "finalize_mode"):
            new_rpt.finalize_mode = None

        db.session.add(new_rpt)
        db.session.commit()

        return jsonify({
            "status": "ok",
            "report_id": new_rpt.id,
            "version_no": new_rpt.version_no,
            "reference_no": new_rpt.reference_no,
            "status_str": new_rpt.status
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.exception("create_version failed")
        return jsonify({"status": "error", "message": str(e)}), 500


@reports_ui_bp.route('/invoice/<int:invoice_id>/generate_pdf', methods=['POST'])
def generate_invoice_pdf_route(invoice_id):
    from .models import GilInvoice, GilInvoiceDocument
    from .reports_docx import generate_invoice_pdf
    from . import db

    invoice = GilInvoice.query.get_or_404(invoice_id)

    payload = request.get_json(silent=True) or {}
    override = bool(payload.get("override", False))

    # --------------------------------
    # Business validation: existing PDF
    # --------------------------------
    has_existing_pdf = False

    if invoice.latest_pdf_path:
        has_existing_pdf = True
    else:
        existing_pdf_doc = (
            GilInvoiceDocument.query
            .filter_by(invoice_id=invoice.invoice_id, document_type='pdf', is_current=True)
            .first()
        )
        if existing_pdf_doc:
            has_existing_pdf = True

    if has_existing_pdf and not override:
        return jsonify({
            "success": False,
            "requires_confirm": True,
            "message": "כבר הופקה חשבונית לתיק הזה. האם לדרוס את הקובץ הקיים?"
        }), 200

    try:
        upload_res = generate_invoice_pdf(invoice)

        invoice.latest_pdf_path = upload_res["dropbox_path"]
        invoice.latest_pdf_filename = upload_res["stored_name"]

        db.session.commit()

        return jsonify({
            "success": True,
            "pdf_path": upload_res["dropbox_path"]
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({
            "success": False,
            "message": str(e)
        }), 500